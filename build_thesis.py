#!/usr/bin/env python3
"""Build FYP thesis DOCX with template-matching formatting and figures."""

import os
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_DIR    = os.path.dirname(os.path.abspath(__file__))
TEMPLATE    = os.path.join(BASE_DIR, "Shema Hugues-25603 (1) (1).docx")
FIGS        = os.path.join(BASE_DIR, "figures")
OUTPUT      = os.path.expanduser("~/Downloads/FYP_Report_FINAL.docx")
TNR         = "Times New Roman"

FIGURE_PATHS = {
    1: os.path.join(FIGS, "Figure 1.png"),
    2: os.path.join(FIGS, "Figure 2.png"),
    3: os.path.join(FIGS, "Figure 3 class diagrams.png"),
    4: os.path.join(FIGS, "Figure 4.png"),
    5: os.path.join(FIGS, "Figure 5.png"),
    6: os.path.join(FIGS, "Figure 6.png"),
    7: os.path.join(FIGS, "Figure 7 student journey .png"),
    8: os.path.join(FIGS, "Figure 8 ERD.png"),
    9: os.path.join(FIGS, "Figure 9.png"),
}

# ── Open template & clear body ────────────────────────────────────────────────
doc = Document(TEMPLATE)
body = doc.element.body

# Save document-level sectPr
sect_pr = body.find(qn('w:sectPr'))
for child in list(body):
    body.remove(child)
if sect_pr is not None:
    body.append(sect_pr)

# Fix margins
s = doc.sections[0]
s.page_width    = Inches(8.5)
s.page_height   = Inches(11.0)
s.left_margin   = Cm(1.905)
s.right_margin  = Cm(1.905)
s.top_margin    = Cm(2.43)
s.bottom_margin = Cm(2.26)

# ── Helpers ───────────────────────────────────────────────────────────────────

def _ls(para, val=1.5):
    pf = para.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = val
    pf.space_before = Pt(0)
    pf.space_after  = Pt(0)

def _run(para, text, size=12, bold=None, italic=None):
    r = para.add_run(text)
    r.font.name = TNR
    r.font.size = Pt(size)
    if bold    is not None: r.font.bold   = bold
    if italic  is not None: r.font.italic = italic
    return r

def body(text, bold=False):
    p = doc.add_paragraph(style='Body Text')
    _ls(p)
    _run(p, text, 12, bold=bold or None)
    return p

def empty():
    p = doc.add_paragraph(style='Body Text')
    _ls(p)
    _run(p, "")
    return p

def pb():
    p = doc.add_paragraph(style='Normal')
    p.add_run().add_break(WD_BREAK.PAGE)

def h(text, level=2):
    style_name = f'Heading {level}'
    p = doc.add_paragraph(style=style_name)
    st = doc.styles[style_name]
    sz = st.font.size.pt if st.font.size else 16
    bold = st.font.bold if st.font.bold is not None else True
    _run(p, text, sz, bold=bold)
    return p

def nc(text, size=14, bold=False):
    """Normal centered paragraph with explicit size."""
    p = doc.add_paragraph(style='Normal')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, text, size, bold=bold or None)
    return p

def ameta(label, value):
    """Abstract metadata line: bold label + normal value."""
    p = doc.add_paragraph(style='Normal')
    _run(p, label, 12, bold=True)
    _run(p, value, 12)
    return p

def sig(text):
    """Signature / date line."""
    parts = text.split(':', 1)
    p = doc.add_paragraph(style='Normal')
    _run(p, parts[0] + ':', 12, bold=True)
    if len(parts) == 2:
        _run(p, parts[1], 12)
    return p

def body_label(label, rest):
    """Body paragraph: bold label colon + normal continuation."""
    p = doc.add_paragraph(style='Body Text')
    _ls(p)
    _run(p, label + ": ", 12, bold=True)
    _run(p, rest, 12)
    return p

def fig(num, caption, img_key=None):
    """Image (if available) + bold centered caption."""
    empty()
    img_path = FIGURE_PATHS.get(img_key or num)
    if img_path and os.path.exists(img_path):
        ip = doc.add_paragraph(style='Normal')
        ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ip.add_run().add_picture(img_path, width=Inches(6.3))
    empty()
    cp = doc.add_paragraph(style='Normal')
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(cp, f"Figure {num}: {caption}", 12, bold=True)
    empty()

def placeholder_fig(num, caption, desc):
    """Caption + italic placeholder description (no image)."""
    empty()
    cp = doc.add_paragraph(style='Normal')
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(cp, f"Figure {num}: {caption}", 12, bold=True)
    dp = doc.add_paragraph(style='Body Text')
    _ls(dp)
    _run(dp, f"[{desc}]", 12, italic=True)
    empty()

def li(text):
    p = doc.add_paragraph(style='List Paragraph')
    _ls(p)
    _run(p, text, 12)
    return p

def _cell(cell, text, bold=False, size=11, italic=False):
    p = cell.paragraphs[0]
    p.clear()
    p.style = doc.styles['Table Paragraph']
    _run(p, text, size, bold=bold or None, italic=italic or None)

def _borders(tbl):
    tpr = tbl.find(qn('w:tblPr'))
    if tpr is None:
        tpr = OxmlElement('w:tblPr')
        tbl.insert(0, tpr)
    tb = OxmlElement('w:tblBorders')
    for name in ('top','left','bottom','right','insideH','insideV'):
        b = OxmlElement(f'w:{name}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), '000000')
        tb.append(b)
    tpr.append(tb)

def tbl(headers, rows, widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    _borders(t._tbl)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    for i, h_text in enumerate(headers):
        _cell(t.rows[0].cells[i], h_text, bold=True)
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            if isinstance(val, tuple):
                _cell(t.rows[ri+1].cells[ci], val[0], bold=val[1])
            else:
                _cell(t.rows[ri+1].cells[ci], val)
    empty()
    return t

def uc_table(pairs):
    """Two-column use-case table from (field, value) pairs."""
    t = doc.add_table(rows=len(pairs), cols=2)
    t.style = 'Table Grid'
    _borders(t._tbl)
    for i, (f, v) in enumerate(pairs):
        _cell(t.rows[i].cells[0], f, bold=True, size=11)
        _cell(t.rows[i].cells[1], v, size=11)
    empty()

# ═══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════════════════════════

nc("Adventist University of Central Africa", 14)
for _ in range(4): empty()
nc("FYP Tracking & Accountability System", 18)
nc("(Final Year Project Tracking & Accountability Platform)", 14)
p4 = doc.add_paragraph(style='Heading 4')
_run(p4, "Case study: AUCA Final Year Project Department", 16)
for _ in range(3): empty()
nc("A Final Year Project Presented in partial fulfillment of the requirements for the degree of "
   "BACHELOR OF SCIENCE IN INFORMATION TECHNOLOGY", 14)
for _ in range(2): empty()
nc("Major in", 14)
for _ in range(2): empty()
nc("Software Engineering", 14)
for _ in range(6): empty()
nc("By MANISHIMWE KWIZERA Jean Luc — June, 2026", 14)

# ═══════════════════════════════════════════════════════════════════════════════
# ABSTRACT
# ═══════════════════════════════════════════════════════════════════════════════
pb()
h("ABSTRACT", 2)
body("A final year project for the Bachelor’s Degree of Science in Information Technology Emphasis in Software Engineering")
body("Adventist University of Central Africa")
empty()
ameta("TITLE: ", "FYP Tracking & Accountability System")
ameta("Name of the Researcher: ", "MANISHIMWE KWIZERA Jean Luc")
ameta("Name of faculty Advisor: ", "ISHIMWE M. Prince")
ameta("Date Completed: ", "June, 2026")
empty(); empty()
body("The FYP Tracking & Accountability System is a web-based platform developed to modernize the management of the final year project journey at the Adventist University of Central Africa. It replaces fragmented manual and semi-digital procedures with a single, auditable system that records every stage a student passes through, from registration to final defense. The platform provides a centralized solution that improves coordination, transparency, and accountability among students, supervisors, examiners, facilitators, and the Head of Department, creating a more organized and evidence-driven academic environment.")
empty()
body("The current process keeps good records at the start of the project journey (registration and case-letter approval) and near the end (book submission and defense scheduling), but has little visibility during prototype review and supervision. These blind spots are where students are most often disadvantaged and where responsibility is assigned by assumption rather than by evidence. Departmental staff face difficulties in tracking supervision meetings, counting prototype and proposal attempts, and proving which notifications were actually sent. The absence of a continuous, timestamped record limits accountability and weakens the resolution of disputes.")
empty()
body("Data were collected through documentation review, direct observation of the departmental workflow, and interviews with academic staff who coordinate final year projects. System design followed an object-oriented approach using UML diagrams. The system was implemented using React with TypeScript on the front end, Spring Boot on the back end, and PostgreSQL for data storage, with Redis for caching and JSON Web Tokens for authentication, ensuring a scalable and maintainable architecture.")
empty()
body("This integrated tracking system centralizes student state management, supervision scheduling, proposal-attempt tracking, panel assignment, notification logging, and audit recording into a single secure platform. It improves operational efficiency, enhances data accuracy, and supports the transition from undocumented manual coordination to a modern, fully digital, and well-coordinated final year project management environment.")

# ═══════════════════════════════════════════════════════════════════════════════
# DECLARATION
# ═══════════════════════════════════════════════════════════════════════════════
pb()
h("DECLARATION", 2)
empty(); empty()
body("I, MANISHIMWE KWIZERA Jean Luc, Student ID Number 26972, a student at the Adventist University of Central Africa in the Faculty of Information Technology, Department of Software Engineering, do hereby declare that this final year project report entitled “FYP Tracking & Accountability System” is entirely the real reflection of my own original work and experience to the best of my knowledge. It has never been either partially or wholly presented in any university or any higher learning institution for the award of a degree or any other qualification.")
empty(); empty()
sig("Signature: ……………………………………….")
empty(); empty()
sig("Date: ………/……………/….………...")

# ═══════════════════════════════════════════════════════════════════════════════
# APPROVAL
# ═══════════════════════════════════════════════════════════════════════════════
pb()
h("APPROVAL", 2)
empty(); empty()
body("I, ISHIMWE M. Prince, hereby certify that student MANISHIMWE KWIZERA Jean Luc, with registration number 26972, in the Faculty of Information Technology, Department of Software Engineering, has completed this Final Year Project report under my supervision and is hereby submitted with my approval.")
empty(); empty()
sig("Signature: ……………………………………………………….")
empty(); empty(); empty()
sig("Date: ………/……………/…………………………………")
for _ in range(10): empty()

# ═══════════════════════════════════════════════════════════════════════════════
# DEDICATION
# ═══════════════════════════════════════════════════════════════════════════════
pb()
h("DEDICATION", 2)
body("To my parents,")
empty()
body("To my beloved sisters and brothers, To my classmates, friends, and relatives,")
body("To my supervisor for his guidance I dedicate this final report")

# ═══════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (static)
# ═══════════════════════════════════════════════════════════════════════════════
pb()
ptoc = doc.add_paragraph(style='Normal')
ptoc.alignment = WD_ALIGN_PARAGRAPH.LEFT
_run(ptoc, "Table of Contents", 16)
empty()
for line in [
    "ABSTRACT",
    "DECLARATION",
    "APPROVAL",
    "DEDICATION",
    "LIST OF FIGURES",
    "LIST OF TABLES",
    "LIST OF ABBREVIATION",
    "ACKNOWLEDGMENTS",
    "CHAPTER 1 GENERAL INTRODUCTION",
    "    Introduction",
    "    Background of the Study",
    "    Statement of Problem",
    "    Choice and Motivation of the Study",
    "    Objectives of the Study",
    "    Scope of the Study",
    "    Methodology and Techniques used in the Study",
    "CHAPTER 2: ANALYSIS OF THE EXISTING SYSTEM",
    "    Introduction",
    "    Description of the Existing System Environment",
    "    Description of the Current System",
    "    Analysis of the Current System",
    "    Proposed Solutions",
    "    System Requirements",
    "CHAPTER 3: REQUIREMENT ANALYSIS AND DESIGN",
    "    Use Case Diagram",
    "    Class Diagram",
    "    Sequence Diagrams",
    "    Activity Diagrams",
    "    Database Schema",
    "    Data Dictionary",
    "CHAPTER 4: IMPLEMENTATION",
    "    Technologies Used",
    "    Presentation of the New System",
    "    Software Testing",
    "    Hardware and Software Requirements",
    "CHAPTER 5: CONCLUSION AND RECOMMENDATIONS",
    "    Summary of the Work",
    "    Achievement of Objectives",
    "    Limitations",
    "    Recommendations",
    "    Conclusion",
]:
    pt = doc.add_paragraph(style='toc 1' if not line.startswith('    ') else 'toc 2')
    _run(pt, line.strip(), 12)

# ═══════════════════════════════════════════════════════════════════════════════
# LIST OF FIGURES
# ═══════════════════════════════════════════════════════════════════════════════
pb()
h("LIST OF FIGURES", 2)
for line in [
    "Figure 1: Current System model",
    "Figure 2: Use case diagram",
    "Figure 3: Class diagram",
    "Figure 4: Sequence Diagram",
    "Figure 5: State transition sequence diagram",
    "Figure 6: User signup and login activity diagram",
    "Figure 7: Student journey activity diagram",
    "Figure 8: Database schema",
    "Figure 9: System architecture",
    "Figure 10: Login Page",
    "Figure 11: Dashboard",
    "Figure 12: User Management Page",
    "Figure 13: Student Import Page",
    "Figure 14: Student Profile and State Timeline",
    "Figure 15: Supervision Page",
    "Figure 16: Panel Assignment Page",
    "Figure 17: Notification and Audit Logs",
]:
    body(line)

# ═══════════════════════════════════════════════════════════════════════════════
# LIST OF TABLES
# ═══════════════════════════════════════════════════════════════════════════════
pb()
h("LIST OF TABLES", 2)
for line in [
    "Table 1: Adding users Use Case Table",
    "Table 2: HOD Use Case Table",
    "Table 3: Backend side Table",
    "Table 4: Sequence diagram Table",
]:
    body(line)

# ═══════════════════════════════════════════════════════════════════════════════
# LIST OF ABBREVIATIONS
# ═══════════════════════════════════════════════════════════════════════════════
pb()
h("LIST OF ABBREVIATION", 1)
empty()
tbl(
    ["Abbreviation", "Meaning"],
    [
        ["API",  "Application Programming Interface"],
        ["AUCA", "Adventist University of Central Africa"],
        ["CORS", "Cross-Origin Resource Sharing"],
        ["CRUD", "Create, Read, Update, Delete"],
        ["CSV",  "Comma-Separated Values"],
        ["DBMS", "Database Management System"],
        ["FYP",  "Final Year Project"],
        ["HOD",  "Head of Department"],
        ["HTTP", "Hypertext Transfer Protocol"],
        ["ICT",  "Information and Communication Technology"],
        ["IT",   "Information Technology"],
        ["JPA",  "Java Persistence API"],
        ["JWT",  "JSON Web Token"],
        ["NCSA", "National Cyber Security Authority"],
        ["NST",  "National Strategy for Transformation"],
        ["ORM",  "Object-Relational Mapping"],
        ["REST", "Representational State Transfer"],
        ["SDLC", "Software Development Life Cycle"],
        ["SQL",  "Structured Query Language"],
        ["UI",   "User Interface"],
        ["UML",  "Unified Modeling Language"],
    ],
    widths=[1.5, 4.5]
)

# ═══════════════════════════════════════════════════════════════════════════════
# ACKNOWLEDGMENTS
# ═══════════════════════════════════════════════════════════════════════════════
pb()
h("ACKNOWLEDGMENTS", 2)
body("First and foremost, I would like to express my sincere gratitude to the Almighty God for His blessings, wisdom, and unwavering guidance throughout the journey of this project. His grace has been my strength and inspiration in every phase of this work, and I am eternally thankful for His divine presence in my life.")
empty()
body("I would like to express my sincere appreciation to AUCA and its entire administration and academic staff for their support and guidance throughout my studies. Special thanks go to the Faculty of Information Technology, particularly the Department of Software Engineering, for their valuable knowledge, resources, and academic support. The high-quality education I received at AUCA has been instrumental in my academic growth and has greatly contributed to the successful completion of this project.")
empty()
body("I wish to extend my heartfelt thanks to my project supervisor, ISHIMWE M. Prince, for his invaluable mentorship, expert guidance, and immense patience. His insightful feedback, constant encouragement, and scholarly advice were crucial in shaping the “FYP Tracking & Accountability System” from a concept into a realized project.")
empty()
body("I am profoundly grateful to my beloved parents for their unconditional love, endless support, and steadfast encouragement. Your belief in my abilities and the countless sacrifices you have made form the bedrock of my achievements. Thank you for always being my pillars of strength.")
empty()
body("I am also thankful to my friends and classmates at AUCA for their camaraderie, collaborative spirit, and moral support throughout our studies. The stimulating discussions and shared experiences have made this academic journey both memorable and rewarding.")
empty()
body("Finally, I would like to sincerely thank everyone who contributed to this work, directly or indirectly. Your support, guidance, and encouragement were invaluable, and I am deeply grateful.")
empty()
body("May God bless you all abundantly")
empty(); empty(); empty()
body("MANISHIMWE KWIZERA Jean Luc")

# ═══════════════════════════════════════════════════════════════════════════════
# CHAPTER 1
# ═══════════════════════════════════════════════════════════════════════════════
pb()
h("CHAPTER 1 GENERAL INTRODUCTION", 2)
h("Introduction", 3)
empty()
body("Advancements in information technology have changed how institutions manage their operations, encouraging the adoption of digital systems that improve efficiency, accuracy, and service delivery. Despite this progress, many academic departments still rely heavily on manual and semi-digital procedures for handling essential activities such as registering students for their final year projects, approving case-study letters, tracking supervision meetings, recording proposal submissions, and assigning examination panels. These paper-based and ad-hoc methods often result in slow processes, difficulties in tracking records, limited transparency, and challenges in maintaining accurate academic and administrative data.")
empty()
body("Departmental coordinators frequently encounter problems when trying to monitor a student’s current stage, verify how many times a prototype or proposal was presented, or confirm whether a required notification was actually delivered. At the same time, students may experience uncertainty about what step comes next, lack timely confirmation after submitting their work, and have limited access to up-to-date information regarding their supervision schedule or panel assignment. Such inefficiencies not only affect the daily administration of final year projects but can also lead to disputes, missed milestones, and reduced trust between students and the department.")
empty()
body("To overcome these challenges, the FYP Tracking & Accountability System is introduced as a digital platform designed to automate and simplify the management of the final year project journey. The system provides a centralized environment where the Head of Department and facilitators can oversee every student’s state efficiently, supervisors can post availability and confirm meetings, examiners can record panel outcomes, and every transition is captured with the actor who performed it and the time it occurred. Students benefit from the ability to follow their own progress online, receive instant email notifications at each milestone, and access the contact details and meeting schedule of their assigned supervisor.")
empty()
body("The adoption of this system is expected to reduce manual workload, improve data accuracy, and strengthen transparency in final year project administration. With real-time access to reliable, timestamped information, departmental staff will be able to make more informed decisions, resolve complaints from a documented record, and respond quickly to operational needs.")

h("Background of the Study", 3)
empty()
body("Final year projects have historically played a central role in higher education, particularly in information technology programmes where they represent the culmination of a student’s academic training and a major source of practical, demonstrable competence. These projects provide opportunities for students to apply theoretical knowledge to real-world problems while also allowing the institution to assess readiness for graduation. Because of their academic and professional importance, effective management of the final year project journey is necessary to maintain order, ensure fair treatment of every student, and support accountable academic decision-making. However, in many institutions, the administration of final year projects still relies heavily on traditional manual systems that were originally designed for smaller cohorts and less complex supervision arrangements.")
empty()
body("Manual final year project management typically involves paper-based registration, handwritten approval of case-study letters, informal scheduling of supervision meetings, and physical or scattered storage of records relating to prototype and proposal attempts. While these practices may appear simple, they often lead to a range of operational challenges. Records can easily be misplaced or contradicted, updating information requires significant time and effort, and verifying the accuracy of data becomes difficult when multiple records are handled manually across several actors. As cohorts grow and the number of students increases, these challenges become more pronounced, leading to delays in coordination, inconsistencies in oversight, and difficulties in tracking each student’s true progress. In addition, the absence of a centralized digital record limits the ability of administrators to generate timely and reliable reports, which are essential for planning, monitoring performance, and resolving disputes.")
empty()
body("The rapid advancement of information and communication technologies (ICT) has introduced new opportunities for improving administrative systems across different sectors. Around the world, organizations and academic institutions are adopting digital platforms to automate routine processes, enhance data accuracy, and improve communication between stakeholders. Electronic management systems allow real-time monitoring of activities, automated record keeping, and faster generation of analytical reports. Such systems not only reduce administrative workload but also increase transparency and accountability by ensuring that information is stored securely and can be accessed whenever needed.")
empty()
body("In Rwanda, the government has demonstrated a strong commitment to promoting digital transformation through national initiatives such as Vision 2050, the National Strategy for Transformation (NST), and the Smart Rwanda Master Plan. These initiatives emphasize the use of technology to enhance service delivery, strengthen governance, and support sustainable development. Significant progress has already been made in sectors such as banking, education, healthcare, and public administration, where digital systems are increasingly being used to manage services and transactions. Despite these advancements, some areas of academic administration, including the management of final year projects, still rely on manual or semi-digital processes that limit efficiency and transparency.")
empty()
body("From the students’ perspective, manual systems can also create inconveniences. Following up on the status of a case letter may require repeated visits to departmental offices, confirmation of supervision meetings may depend on informal messages, and accessing updated information about prototype outcomes, proposal decisions, or panel schedules may not always be straightforward. Such challenges can affect a student’s ability to plan their work, reduce confidence in the fairness of the process, and create unnecessary delays. As cohorts continue to expand and the demand for accountable academic services increases, the need for a more reliable and technology-driven management approach becomes increasingly evident.")
empty()
body("It is within this context that the FYP Tracking & Accountability System is proposed as a comprehensive digital solution aimed at improving the management of the final year project journey. The system is designed to provide an integrated platform where the Head of Department and facilitators can oversee every stage of the workflow, supervisors can post availability and confirm meetings, examiners can record pre-defense and defense outcomes, and students can follow their own progress and receive timely notifications. Each transition is recorded with the actor and a timestamp, and every email the system sends is stored in a notification log so that claims of “I was never told” can be checked against what the system actually sent.")
empty()
body("By replacing fragmented manual procedures with an automated and centralized system, the FYP Tracking & Accountability System is expected to significantly enhance operational efficiency, data accuracy, and transparency in final year project administration. The system will help reduce paperwork, minimize human errors, and ensure that reliable, timestamped information is readily available for students, supervisors, examiners, facilitators, and the Head of Department alike.")

h("Statement of Problem", 3)
empty()
body("Current final year project management methods are largely inefficient, relying on manual and semi-digital processes for registration, supervision tracking, prototype and proposal recording, and panel coordination, leading to operational delays, accountability gaps, and poor resource utilization. The department keeps adequate records at the start of the journey and near the end, but has little visibility during prototype review and supervision, which are precisely the stages where students are most often disadvantaged and where responsibility is today assigned by assumption rather than by evidence. These outdated practices are inadequate for modern academic demands, especially in growing programmes where efficiency and transparency are crucial. The main challenges that characterize the existing approach include:")
empty()
body_label("Supervision blind spot", "There is no systematic record of supervisor availability, confirmed meeting dates and times, attendance, or book sign-off, making it impossible to verify whether supervision actually took place as required.")
empty()
body_label("Lack of real-time state visibility", "Departmental staff lack access to up-to-date information on each student’s current stage and history, hindering effective coordination, follow-up, and decision-making.")
empty()
body_label("Untracked prototype and proposal attempts", "The number of times a student presented a prototype, or submitted and had a proposal rejected, is not reliably recorded with reasons, so a student’s true history cannot be proven when a dispute arises.")
empty()
body_label("Administrative inefficiency", "The reliance on manual processes creates significant administrative overhead, delays in moving students between stages, and difficulties in maintaining accurate records across multiple actors.")
empty()
body_label("Poor communication and notification evidence", "Students often miss important updates regarding approvals, supervisor assignment, or panel schedules, and the department cannot prove which messages were sent, leading to disputes that cannot be resolved from a record.")
empty()
body_label("Difficulty in monitoring and reporting", "Without a digital system, tracking how long students spend at each stage, which panels are assigned, and which actions were taken by whom requires extensive manual effort, making it hard for the department to identify bottlenecks, ensure accountability, or generate timely reports for decision-making.")

h("Choice and Motivation of the Study", 3)
empty()
body("The motivation behind the FYP Tracking & Accountability System is to provide a comprehensive digital platform that streamlines final year project management by integrating real-time state tracking, supervision scheduling, proposal-attempt recording, and panel coordination into one auditable timeline. By doing so, the FYP Tracking & Accountability System aims to reduce administrative burdens, enhance transparency, and improve resource utilization, addressing the fundamental challenges of traditional final year project management while promoting an accountable and evidence-driven academic environment.")
empty()
body_label("To AUCA", "Developing the FYP Tracking & Accountability System provides a practical opportunity to apply the knowledge and skills I have gained during my Bachelor of Science in Information Technology studies at the Adventist University of Central Africa (AUCA). This project integrates system design, full-stack web development, and secure data management, aligning with AUCA’s mission of nurturing innovative professionals who leverage technology to solve real-world problems and promote efficient, community-centered digital transformation.")
empty()
body_label("To the AUCA FYP Department", "The FYP Tracking & Accountability System directly supports the department’s effort to modernize academic administration and ensure fair, accountable treatment of every final year student. By providing a centralized digital platform for state tracking, supervision, proposal recording, panel assignment, communication, and audit, the system improves operational efficiency, transparency, and accountability across the entire project journey.")
empty()
body("The platform also aligns with Rwanda’s Vision 2050 and the National Strategy for Transformation (NST) by promoting inclusive digital innovation, data-driven management, and stronger institutional governance. Furthermore, it contributes to the Smart Rwanda Master Plan by leveraging ICT to strengthen academic service delivery, streamline departmental operations, and improve the overall management of the final year project process.")
empty()
body_label("To the Student", "This project represents both an academic and personal commitment to applying technology for practical impact. As a student specializing in Software Engineering, I was inspired by the real challenges faced in managing the final year project journey efficiently and fairly. The FYP Tracking & Accountability System enabled me to transform theoretical knowledge into a real-world solution that promotes accountability, streamlines academic processes, and supports Rwanda’s innovation and development agenda.")

h("Objectives of the Study", 3)
h("General Objective", 5)
body("To develop a digital platform that streamlines final year project management by automating state tracking, supervision scheduling, proposal-attempt recording, and panel coordination through real-time data integration. The system will provide the department with tools for efficient operations, transparency, and accountability, using a relational database and caching to deliver reliable, timestamped records. Automated notifications and milestone updates will guide students, supervisors, and examiners, reducing delays and disputes. Accessible through a responsive web application, the platform aims to lower administrative costs, improve student satisfaction, and create a more organized and accountable final year project process.")
h("Specific Objectives:", 5)
for obj in [
    "To develop a secure, web-based, and responsive platform that enables the department, supervisors, examiners, and students to manage registration, state transitions, supervision, and panels in real time, reducing reliance on manual and paper-based processes.",
    "To design an intuitive and user-friendly interface that accommodates all roles — student, supervisor, facilitator, HOD, examiner, and superadmin — ensuring ease of use, accessibility, and wider adoption across different users.",
    "To build a robust database system for securely storing and managing student profiles, state transitions, supervision records, proposal attempts, panel assignments, notification logs, and audit logs efficiently.",
    "To enforce a clear state machine and role-based business rules that govern valid transitions, count prototype and proposal attempts, and prevent invalid or unauthorized actions across the project journey.",
    "To implement an automated communication and notification system that alerts students, supervisors, and examiners about approvals, supervisor assignment, meetings, and panel schedules, and records every message in a notification log.",
    "To provide oversight and reporting tools that allow the department to monitor each student’s state and timeline, track panel assignments, and review immutable audit and notification records to support data-driven decisions.",
    "To ensure data security, privacy, and system reliability by implementing authentication, authorization, and encryption mechanisms to protect sensitive student information in line with Rwanda’s personal data protection law.",
]:
    li(obj)

h("Scope of the Study", 3)
empty()
body("The FYP Tracking & Accountability System will focus on the final year project administration of an academic department, providing a comprehensive digital management platform that integrates state tracking, supervision scheduling, proposal-attempt recording, panel assignment, and communication features. By optimizing existing departmental operations, the platform will help the Head of Department and facilitators coordinate the journey efficiently and make evidence-based decisions regarding each student’s progress. In addition to providing operational visibility, the system will track prototype attempts, proposal rejection reasons, supervision events, and panel outcomes, promoting transparent and accountable project management without requiring significant infrastructure changes.")
empty()
body("Rather than replacing the institution’s existing registrar or prototype-review systems, the primary goal of the FYP Tracking & Accountability System is to enhance operational efficiency and accountability using digital solutions. By leveraging real-time data, a clear state machine, and automated workflows, the platform will support the department in adopting modern management practices, ultimately reducing administrative burdens and improving student satisfaction. This approach ensures accessibility and practicality for departments managing roughly 500 to 1000 students per cycle while contributing to long-term institutional sustainability.")
empty()
body("The study will concentrate on developing the core software platform and its essential modules, including:")
for item in [
    "State tracking and per-student timeline management",
    "Supervision module: availability, supervisor-confirmed meetings, and book sign-off",
    "Proposal-attempt tracking with rejection reasons and attempt limits",
    "Panel assignment for pre-defense and defense, with outcome recording",
    "Automated milestone notifications and a notification log",
    "Role-based access control, authentication, and an immutable audit log",
]:
    li(item)
empty()
body("The implementation will utilize existing technological infrastructure commonly available in academic environments, focusing on web accessibility to ensure broad usability across the department’s roles.")

h("Methodology and Techniques used in the Study", 3)
empty()
body("The research methodology for the FYP Tracking & Accountability System will adopt a systematic approach to analyzing and improving final year project management processes. This will include both qualitative and quantitative techniques, such as literature review, direct observation of departmental operations, interviews with academic staff, and case studies of existing digital management systems. These methods will help gain a comprehensive understanding of current project management practices, identify operational inefficiencies, and develop solutions to optimize departmental operations through digital transformation.")
h("Documentation", 6)
body("Documentation is a research method that involves the systematic collection, review, and analysis of existing records, reports, and written materials relevant to the study topic. It allows researchers to understand background information, identify operational challenges, and gather evidence to support research objectives.")
empty()
body("In this study, I used documentation review to examine existing records, forms, and guidelines used by the department to manage final year projects. This included analyzing student registration records, case-letter approval forms, supervision arrangements, prototype and proposal records, and the requirements brief prepared for the system, in order to understand the journey, the actors, and the data each stage must capture.")
h("Observation", 5)
body("Observation is a qualitative research method that involves directly monitoring and recording behaviors, processes, and interactions in a real-world setting without interference. In the context of the FYP Tracking & Accountability System, observation was used to study how the department currently manages registration, supervision, prototype and proposal review, and panel coordination. This method allowed the collection of objective data on administrative workflows, decision-making patterns, and operational challenges affecting the efficiency and fairness of final year project management.")
h("Interview", 5)
body("Interviews are a qualitative research method used to gather in-depth insights from individuals with relevant knowledge and experience. They involve structured, semi-structured, or open-ended discussions to explore perspectives, challenges, and potential solutions related to a specific topic. In the context of the FYP Tracking & Accountability System, interviews were conducted with departmental staff who coordinate final year projects.")
empty()
body("I, MANISHIMWE KWIZERA Jean Luc, a finalist in the Faculty of Information Technology, Department of Software Engineering at the Adventist University of Central Africa (AUCA), conducted this research to design and develop the FYP Tracking & Accountability System, specifically addressing inefficiencies in state tracking, supervision, proposal recording, and panel coordination.")
empty()
body_label("Question 1", "What challenges do you currently face in managing the final year project journey? Answer: The main challenges include limited visibility during supervision, difficulty in tracking how many times a student presented a prototype or resubmitted a proposal, lack of a single view of each student’s stage, and no reliable proof of which notifications were sent. These challenges often lead to coordination delays, disputes over what happened, and unfair outcomes for some students.")
empty()
body_label("Question 2", "How effective are the current tools and methods used for managing final year projects? Answer: The current tools and methods are mostly manual or scattered across forms, spreadsheets, and informal messages. While they have been used for a long time, they are inefficient, prone to errors, and unable to provide a real-time, timestamped record of supervision, attempts, and communications.")
empty()
body_label("Question 3", "What improvements are needed to enhance final year project management? Answer: There is a strong need for a digital solution that records every state a student passes through, captures supervision and panel events with timestamps, counts prototype and proposal attempts, and stores every notification sent. Improving transparency, reducing manual workload, and enabling accountability are essential.")
empty()
body_label("Question 4", "What features do you expect from a new final year project management system? Answer: The expected features include real-time state tracking and per-student timelines, supervisor availability and confirmed meetings, proposal-attempt tracking with reasons, examiner panel assignment, automated milestone emails, and immutable audit and notification logs.")
empty()
body_label("Question 5", "How would a digital final year project management system benefit departmental operations? Answer: A digital system would improve operational efficiency, increase transparency, reduce errors, and support better decision-making and dispute resolution through real-time, timestamped data and reliable records.")

h("Expected Results", 5)
empty()
body("The FYP Tracking & Accountability System is expected to effectively address the challenges identified in the objectives of this project. After successful implementation,")
empty()
body_label("Improved State Visibility and Coordination", "The system will enable real-time tracking of every student’s state and timeline. This will reduce confusion about progress, eliminate inconsistent records, and ensure that the department can coordinate the journey efficiently from a single source of truth.")
empty()
body_label("Closed Supervision Blind Spot", "By digitizing supervisor availability, confirmed meetings, attendance, and book sign-off, the system will make supervision verifiable. All supervision events will be timestamped, turning the record into evidence neither side can later deny.")
empty()
body_label("Provable Attempt and Rejection Trail", "Integration of prototype-attempt visibility and proposal-attempt tracking will record how many times a student presented and the reasons for any rejection, making a student’s history provable rather than anecdotal.")
empty()
body_label("Faster Communication and Notification Evidence", "The system will facilitate instant communication through automated milestone emails, and store every message in a notification log with delivery status. This will improve information flow and allow the department to prove what was sent and when.")
empty()
body_label("Stronger Accountability and Decision Making", "Immutable audit logs and per-student timelines will provide the department with reliable insight into who acted and when, supporting informed planning, fair outcomes, and dispute resolution.")
empty()
body_label("Enhanced Student Satisfaction and Process Efficiency", "By clarifying each student’s next step, confirming supervision, and recording every decision, the system will reduce administrative burdens, improve the student experience, and contribute to a more organized and accountable final year project process.")

h("Organization of the Work", 5)
empty()
body("This research study comprises five chapters, each serving a distinct purpose in the development and analysis of the FYP Tracking & Accountability System:")
empty()
body_label("Chapter One", "General Introduction, provides a general overview of the project, including the background of the study, problem statement, motivation, objectives (general and specific), scope of the system, methodology, requirements collection techniques, and expected results.")
empty()
body_label("Chapter Two", "Analyzes the existing manual final year project management approach, identifies its weaknesses, and defines the requirements of the proposed FYP Tracking & Accountability System. It includes functional and non-functional requirements, the current process model, the problems of the existing system, and the expected benefits.")
empty()
body_label("Chapter Three", "Requirement analysis and design of the new system, presents the system design, including the high-level system architecture, module descriptions, and overall solution design. It also covers process models, the database schema, class diagrams, use-case diagrams, sequence diagrams, and activity diagrams.")
empty()
body_label("Chapter Four", "Implementation of the FYP Tracking & Accountability System, focuses on the implementation phase of the system. It describes the tools and technologies used, system development, coding practices, database implementation, and testing procedures applied to ensure system reliability and performance.")
empty()
body_label("Chapter Five", "Conclusion and recommendations, summarizes the project, highlights achievements, discusses limitations encountered during development, and proposes possible future enhancements to improve the FYP Tracking & Accountability System.")

# ═══════════════════════════════════════════════════════════════════════════════
# CHAPTER 2
# ═══════════════════════════════════════════════════════════════════════════════
pb()
h("CHAPTER 2", 2)
pch2 = doc.add_paragraph(style='Normal')
pch2.alignment = WD_ALIGN_PARAGRAPH.CENTER
_run(pch2, "ANALYSIS OF THE EXISTING SYSTEM", 16, bold=True)

h("Introduction", 3)
empty()
body("This chapter provides a comprehensive analysis of the existing final year project management approach and lays the foundation for the development of the FYP Tracking & Accountability System. The primary objective of this section is to examine current departmental practices, identify operational challenges, and demonstrate how the proposed system can address these gaps to enhance efficiency, transparency, and overall accountability.")
empty()
body("The FYP Tracking & Accountability System is designed to streamline the administration of the final year project journey by providing a centralized digital platform for students, supervisors, examiners, facilitators, and the Head of Department. Its core functionalities include real-time tracking of each student’s state, management of supervisor availability and confirmed meetings, recording of prototype and proposal attempts, assignment of pre-defense and defense panels, generation of operational records, and automated communication tools that keep all stakeholders informed. By leveraging these features, the system aims to reduce human error, minimize operational delays, and create a more transparent and accountable academic environment.")
empty()
body("Current final year project management practices, particularly in traditional settings, often rely on manual processes, such as paper-based registration and approval, informal scheduling of supervision meetings, and scattered tracking of attempts and decisions. These methods are prone to inefficiencies, including lost or contradictory records, unverifiable supervision, and difficulties in monitoring student progress. Students also face challenges such as unclear next steps, lack of timely notifications, and limited access to their own records, which can result in misunderstandings and reduced trust in departmental administration. From an accountability perspective, these inefficiencies contribute to unresolved disputes and hinder fair treatment of every student.")
empty()
body("The FYP Tracking & Accountability System proposes solutions that address these limitations. By centralizing student information, automating state transitions and notifications, and providing real-time timelines, the system enhances operational efficiency, promotes accountability, and ensures accurate, timestamped records throughout the project journey.")

h("Description of the Existing System Environment Historical Background", 5)
body("The Adventist University of Central Africa (AUCA) is a private institution of higher learning based in Kigali, Rwanda, committed to nurturing competent, ethical, and innovative professionals. Within its Faculty of Information Technology, the department responsible for final year projects coordinates the journey that every graduating student must complete, from initial registration through supervision, proposal review, and the final defense. This journey is a core academic requirement and represents the point at which students demonstrate their ability to apply the knowledge and skills acquired throughout their studies.")
empty()
body("Over the years, the department has managed an increasing number of final year students each cycle, with multiple supervisors, examiners, and coordinating staff involved at different stages. The process spans several distinct phases — registration, case-letter approval, prototype review, proposal review, supervision, pre-defense, and defense — each requiring records to be captured and decisions to be communicated. As the cohort has grown, the coordination of these phases has become increasingly demanding, particularly in the supervision and prototype-review stages where oversight has traditionally been weakest.")
empty()
body("Among the department’s priorities is ensuring that every student is treated fairly and that academic decisions can be justified from a reliable record. The management of the final year project journey therefore represents a significant administrative responsibility, positioning the department as a key driver of academic quality and accountability within the faculty.")

h("Mission", 5)
body("The mission of the AUCA Final Year Project Department is to coordinate and oversee the final year project journey in a fair, transparent, and efficient manner, ensuring that every student receives proper supervision and timely communication, and that all academic decisions are supported by reliable, timestamped records.")

h("Vision", 5)
body("The vision of the AUCA Final Year Project Department is to become a model of accountable and technology-driven academic administration, where the entire final year project journey is managed through a single, auditable platform that strengthens transparency, supports fair outcomes, and contributes to national development goals such as Vision 2050 through quality education and digital innovation.")

h("Description of the current System", 3)
empty()
body("At present, the final year project management process within the department operates largely manually or through scattered semi-digital tools. Registration lists, case-letter approvals, supervision arrangements, and records of prototype and proposal attempts are handled using paper forms, spreadsheets, and informal communication. Departmental staff assign supervisors and approve milestones in person or by message, while students rely on visiting offices or contacting staff directly to learn the status of their work.")
empty()
body("Information about approvals, supervisor assignment, meeting schedules, or panel dates is usually shared through phone calls, in-person notices, or informal group messages. Students who wish to confirm their stage or follow up on a decision must contact staff directly, often resulting in delays and uncertainty. After a decision is made, records are updated manually, and any report summarizing student states, attempts, or supervision is prepared by hand.")
empty()
body("This manual approach creates multiple challenges. Coordination between the Head of Department, facilitators, supervisors, examiners, and students relies heavily on phone calls or in-person follow-ups, which are time-consuming and prone to errors. There is no centralized system to monitor each student’s state, track supervision events, or provide real-time updates, making it difficult to manage the journey efficiently or ensure timely communication with students.")
empty()
body("Furthermore, manual record-keeping makes it difficult to maintain accurate data over time. Lost or contradictory documents, inconsistent records, and delayed consolidation of attempts and supervision records are common problems. As a result, staff struggle to generate reliable reports or to prove what actually happened when a dispute arises.")
empty()
body("The current final year project management process faces several critical challenges that limit efficiency, transparency, and service quality:")
empty()
body_label("Manual State and Milestone Management", "Staff track student stages and approve milestones on paper or in spreadsheets. This process is slow, prone to errors, and difficult to scale as the number of students increases. Students may face delays in obtaining confirmation of their progress.")
empty()
body_label("Delayed Communication and Student Updates", "Announcements about approvals, supervisor assignment, or panel dates are shared through notices, meetings, or messages. This informal communication often leads to missed notifications or inconsistent information, reducing student confidence and participation.")
empty()
body_label("Lack of Real-Time Tracking", "There is no system to monitor student state or supervision events in real time. Staff rely on repeated in-person checks and messages, while students have limited visibility into their own progress.")
empty()
body_label("Poor Data and Document Management", "Records of students, attempts, supervision, and decisions are stored physically or in scattered files, making them vulnerable to loss, duplication, or contradiction. This compromises the accuracy and reliability of departmental data.")
empty()
body_label("Inefficient Coordination Among Stakeholders", "Communication between staff, supervisors, examiners, and students is informal and inconsistent. Without a structured system, follow-ups on overdue supervision, pending approvals, or panel assignment are slow and error-prone.")
empty()
body_label("Limited Reporting and Accountability", "Records of states, attempts, supervision, and communications are prepared manually, lacking depth and an audit trail. Decision-making and dispute resolution are hindered by the absence of a centralized system capable of tracking who acted and when.")

h("Analysis of the Current System", 3)
empty()
body("At the AUCA Final Year Project Department, operations such as registration, case-letter approval, supervision, prototype and proposal review, panel coordination, and communication are still largely managed through manual or semi-digital processes. These fragmented approaches limit operational efficiency and make it difficult to manage a large cohort effectively. The absence of a centralized digital system results in inconsistent records, duplication of data, and delays in accessing accurate information.")
empty()
body("Decision-making within the department is often slowed by the lack of real-time data on student state, supervision compliance, and attempt history. Generating operational records requires manual compilation, which is time-consuming and prone to errors. This limits the department’s ability to monitor progress, optimize coordination, and plan strategically.")
empty()
body("Accountability also faces challenges due to non-integrated record-keeping. Without timestamped supervision events, attempt tracking, and notification logging, monitoring what actually happened becomes difficult, increasing the risk of delays, disputes, and unfair outcomes. In addition, communication between the department and students relies heavily on verbal notices or informal messages, which are not always timely or reliable. This communication gap contributes to missed milestones, misunderstandings, and reduced transparency.")
empty()
body("Furthermore, the lack of an integrated dashboard makes it difficult for staff to oversee the journey holistically, while students have limited access to their own state and supervision information. These challenges reduce accountability and weaken trust between the department and its students.")
empty()
body("The key challenges faced by the existing system at the department include:")
empty()
body_label("Delayed Decision-Making", "Absence of real-time, centralized data on student state, supervision, and attempt history.")
empty()
body_label("Inefficient State and Supervision Management", "Manual tracking and coordination leading to errors and unverifiable supervision.")
empty()
body_label("Weak Accountability", "Lack of timestamped events, attempt tracking, and an immutable audit trail.")
empty()
body_label("Poor Communication", "Ineffective dissemination of approvals, supervisor-assignment notices, and panel updates, with no proof of what was sent.")
empty()
body_label("Limited Transparency and Reporting", "Inadequate tools for monitoring progress and supporting fair, evidence-based decisions.")

h("Modelling of the current System", 5)
body("The current final year project management process operates primarily through manual procedures. Departmental staff rely on paper-based records and informal coordination to register students, approve milestones, arrange supervision, and assign panels. State tracking is often inconsistent, leading to uncertainty about where a student stands. Supervision is arranged informally, making attendance and sign-off difficult to verify. Communication with students is largely informal, causing delays in updates and notifications. Overall, the existing system lacks real-time visibility, automated record-keeping, and centralized data management, resulting in operational inefficiencies, limited transparency, and challenges in decision-making.")
empty()
fig(1, "Current System model", 1)
body("The current process of managing the final year project journey in the department operates entirely manually. The workflow can be described in several sequential steps:")
empty()

h("Step 1: Student Registration", 6)
body("The department, through the Head of Department, compiles the list of students who have registered for their final year project, together with their contact details. This list is maintained manually or in a spreadsheet and serves as the primary record of who is participating in the journey for the cycle.")
h("Step 2: Case-Study Letter Submission", 6)
body("Each student prepares and submits a case-study letter and the accompanying form describing their intended project and organisation. These submissions are made physically or by message and recorded by departmental staff. The paper forms serve as the primary record of the student’s intent to begin the journey.")
h("Step 3: Case-Letter Review and Approval", 6)
body("The Head of Department reviews the submitted case-study letters to confirm eligibility and suitability, then approves them and issues the preliminary requirements. Incomplete or unclear submissions require follow-ups through phone calls or in-person visits, delaying the approval process and the student’s progression to prototype review.")
h("Step 4: Prototype Review", 6)
body("Approved students proceed to prototype review, where they present their prototype and may be asked to refine and re-present it. The number of presentations before a prototype is granted is tracked informally, if at all. Because this stage is handled outside any central record, the department has limited visibility into how many attempts each student made and the decisions taken.")
h("Step 5: Proposal Submission and Review", 6)
body("Once a prototype is granted, the student submits the proposal book. Reviewers accept or reject the proposal, and rejected proposals must be revised and resubmitted. The reasons for rejection and the number of attempts are recorded manually and inconsistently, making it difficult to prove a student’s submission history when a dispute arises.")
h("Step 6: Supervisor Assignment", 6)
body("After the proposal is accepted, the department assigns a supervisor to the student. Notification of the assignment is communicated manually through phone calls or messages, often resulting in delays or missed messages. The student and supervisor then coordinate their work informally.")
h("Step 7: Supervision Meetings", 6)
body("The student meets the assigned supervisor for guidance throughout the project. Supervisor availability, meeting dates and times, attendance, and book sign-off are arranged and recorded informally, with no central, timestamped record. This is the stage where oversight is weakest and where supervision cannot be reliably verified.")
h("Step 8: Pre-Defense and Defense Panel", 6)
body("Once the book is signed off and submitted, the facilitator arranges the pre-defense and defense panels by assigning eligible examiners. Schedules and assignments are communicated manually. Examiners record the pre-defense and defense outcomes on paper, and the results are filed manually, with limited traceability.")
h("Step 9: Reporting and Record-Keeping", 6)
body("Departmental staff summarize state, supervision, attempt, and panel records, preparing manual reports for review and filing. The lack of a centralized system limits the ability to track a student’s full history, generate an audit trail, or monitor progress across the cohort efficiently.")

h("Problem of the Existing System", 5)
body("Current final year project management in many departments relies heavily on manual processes, including paper-based records, informal supervision arrangements, scattered attempt tracking, and informal communication methods such as phone calls or group messages. These practices create inefficiencies, delays, and weak coordination between departmental staff and students. Using the PIECES framework, the main challenges can be summarized as follows:")
empty()
h("Performance", 6)
body("Manual registration, milestone approval, and state tracking slow operations and often lead to confusion or inconsistent records of student progress.")
empty()
body("Supervision arrangement and record-keeping are done informally, increasing the chance of unverifiable meetings, missed sign-offs, and delays in progression.")
h("Information", 6)
body("Student, attempt, and supervision data are inconsistently recorded, with no standardized digital forms or automated validation.")
empty()
body("Records of state, supervision, and panel outcomes are manually prepared, offering no real-time updates or per-student timelines.")
empty()
body("Data is scattered across paper files, spreadsheets, and informal messages, increasing the risk of errors, duplication, or loss of records.")
h("Economics", 6)
body("Manual processes increase administrative costs due to paperwork, in-person coordination, and repeated corrections.")
empty()
body("Staff spend significant time verifying supervision, reconstructing attempt history, or resolving disputes, reducing productivity and departmental efficiency.")
empty()
body("Inefficient coordination leads to wasted effort and higher operational overhead.")
h("Control", 6)
body("Departmental staff have limited visibility into daily operations due to lack of real-time tracking.")
empty()
body("There is no audit trail for state transitions, supervision events, or approvals, complicating accountability and the resolution of disputes.")
h("Efficiency", 6)
body("Coordination between students and staff is slow, with frequent rework and duplicated communication.")
empty()
body("Manual oversight limits the ability to respond quickly to pending milestones, overdue supervision, or panel scheduling needs.")
h("Service", 6)
body("Students often receive delayed or inconsistent updates on approvals, supervisor assignment, or panel schedules.")
empty()
body("Communication relies on multiple informal channels, leading to confusion, missed notifications, and an inability to prove what was sent.")

h("Proposed Solutions", 3)
empty()
body("The proposed solution to address the inefficiencies in current final year project management practices is the implementation of the FYP Tracking & Accountability System, an integrated, data-driven digital management platform.")
empty()
body("Key features and benefits of the proposed platform include:")
empty()
body_label("Real-time State Management", "The platform will provide live tracking of each student’s state and a per-student timeline, enabling clear visibility of progress and the next required step.")
empty()
body_label("User-friendly Interface", "An intuitive dashboard that allows departmental staff, supervisors, examiners, and students, regardless of technical expertise, to easily manage their respective tasks and view relevant information.")
empty()
body_label("Process Automation", "Automation of state transitions, attempt counting, supervisor assignment, and panel coordination to reduce administrative workload and minimize human error.")
empty()
body_label("Accountability and Audit Trail", "An immutable audit log and per-student timeline that record who acted and when, turning the system into reliable evidence for dispute resolution.")
empty()
body_label("Supervision Tracking", "Supervisor availability, confirmed meeting dates and times, attendance, and book sign-off, all captured and timestamped to close the supervision blind spot.")
empty()
body_label("Attempt and Rejection Tracking", "Recording of prototype attempts and proposal submissions with rejection reasons and attempt limits, making a student’s history provable.")
empty()
body_label("Automated Communication System", "Instant milestone emails for approvals, supervisor assignment, meetings, and panel schedules, with a notification log that proves what was sent and its delivery status.")
empty()
body_label("Decision Support Tools", "Oversight views that help departmental staff make data-driven decisions regarding coordination, supervision, and panel management.")

h("System Requirements", 3)
empty()
for req in [
    ("REQ 1", "The system shall allow the Head of Department, facilitators, supervisors, examiners, and students to be provisioned with accounts and to access the platform through a secure, user-friendly interface, with student identity and contact data imported from the registrar’s list."),
    ("REQ 2", "The system shall allow users to log in securely and authenticate their identity using encrypted credentials and token-based sessions."),
    ("REQ 3", "The platform shall implement role-based access control, granting different permissions and access levels to Student, Supervisor, Facilitator, HOD, Examiner, and Superadmin to ensure data security and operational control."),
    ("REQ 4", "The system shall record every state a student passes through, with the actor who performed the transition and a timestamp, and present a per-student timeline of the entire journey."),
    ("REQ 5", "The system shall enforce a defined state machine, allowing only valid transitions, counting prototype re-presentations and proposal attempts, and rejecting invalid or unauthorized actions."),
    ("REQ 6", "The system shall support supervision management, allowing supervisors to post weekly availability, confirm meeting dates and times, record attendance, and sign off the book."),
    ("REQ 7", "The system shall track proposal submissions, recording each attempt with its accept or reject outcome, the rejection reason, and the attempt number, and enforcing the maximum-attempt limit."),
    ("REQ 8", "The system shall allow the facilitator to assign pre-defense and defense examiners to a student, keeping panel assignments independent of supervision, and shall record the panel outcomes."),
    ("REQ 9", "The system shall integrate with the external prototype-review system, read-only via an API queried by student ID, to display the schedule, the number of presentations, and the decisions, caching the last result for availability during disputes."),
    ("REQ 10", "The platform shall include a communication system that sends automated milestone emails — including the supervisor-assignment email carrying the supervisor’s name and a one-click WhatsApp group link — to students, supervisors, and examiners."),
    ("REQ 11", "The system shall maintain a notification log recording every email sent with the recipient, time, contents, and delivery status, and an immutable audit log recording who changed what and when."),
    ("REQ 12", "The platform shall provide oversight views and per-student timelines that allow the department to monitor each student’s state, supervision, attempts, and panel assignments for coordination and decision-making."),
]:
    body_label(req[0], req[1])
    empty()

h("Non-Functional Requirements", 5)
empty()
for req in [
    ("REQ 13", "The system should efficiently process state transitions, supervision records, and notifications within a maximum response time of 5 seconds under normal operating conditions."),
    ("REQ 14", "The system should handle multiple simultaneous users, including staff, supervisors, examiners, and students, without performance degradation during peak periods."),
    ("REQ 15", "The platform shall be scalable to support roughly 500 to 1000 students and 10 to 50 supervisors per cycle through a relational database with appropriate indexing and a caching layer."),
    ("REQ 16", "The system shall maintain high availability, with data backup and recovery mechanisms to ensure uninterrupted access and data preservation."),
    ("REQ 17", "The platform shall comply with standard data protection and cybersecurity measures, including authentication, authorization, and encryption, to safeguard student and departmental data from unauthorized access, in line with Rwanda’s personal data protection law (Law N° 058/2021)."),
    ("REQ 18", "Sensitive information, such as student profiles, contact details, and records, must be encrypted both in transit and at rest, with personal data retained for one year and then removed or anonymised."),
    ("REQ 19", "The system shall be compatible with all major browsers (e.g., Chrome, Firefox, Edge) and accessible across mobile devices, tablets, and desktops, maintaining consistent functionality and layout."),
    ("REQ 20", "The platform shall be developed with a modular architecture and well-documented components to simplify maintenance, troubleshooting, updates, and future feature integration, including the swap-in of the real prototype-review API behind a thin adapter."),
    ("REQ 21", "The system shall ensure data accuracy and integrity by validating user inputs, maintaining immutable audit and notification logs of system activities, and supporting secure backups of student and transaction data."),
    ("REQ 22", "The system shall provide an intuitive and responsive user interface for all users, including students, supervisors, examiners, facilitators, and the Head of Department, supporting web and mobile-responsive access."),
    ("REQ 23", "Integrated help resources and guidance shall be available to support new users, ensuring efficient adoption of the platform."),
]:
    body_label(req[0], req[1])
    empty()

# ═══════════════════════════════════════════════════════════════════════════════
# CHAPTER 3
# ═══════════════════════════════════════════════════════════════════════════════
pb()
h("CHAPTER 3", 2)
pch3 = doc.add_paragraph(style='Heading 2')
_run(pch3, "REQUIREMENT ANALYSIS AND DESIGN OF THE NEW SYSTEM", 16, bold=True)

h("Introduction", 3)
empty()
body("Developing an effective, data-driven solution to address operational inefficiencies requires a comprehensive approach to both system analysis and design. This process can be likened to building a strong foundation, which is essential for creating a resilient and functional framework. System analysis and system design serve as the fundamental pillars of the system development life cycle, guiding the project from understanding needs to delivering a working solution.")
empty()
body("System analysis involves systematically gathering and evaluating user requirements, identifying challenges, and breaking down the system into core components. The primary goal is to fully understand the system’s objectives and ensure that the proposed solution effectively resolves identified inefficiencies in final year project management. By thoroughly analyzing these challenges, the system’s overall performance and functionality are enhanced, ensuring all components work seamlessly together toward achieving the platform’s goals. This phase also helps prioritize features based on user needs and technical feasibility, laying the groundwork for successful implementation.")
empty()
body("Conversely, system design focuses on defining the structure, architecture, components, and interactions necessary to meet the system requirements identified during analysis. It builds on the insights gathered to fill gaps and address unmet needs, producing detailed specifications that clarify both functional and operational aspects of the FYP Tracking & Accountability System. The main objective of system design is to establish clear methods and strategies to achieve the desired system outcomes. A well-executed design phase ensures the platform is scalable, maintainable, and adaptable to evolving technologies and user demands, ultimately leading to a more robust and user-centric final year project management system.")

h("Description of the New System", 3)
empty()
body("The FYP Tracking & Accountability System is a comprehensive digital project-management platform designed to address operational inefficiencies in traditional final year project administration through a data-driven, user-centered approach. By integrating state tracking, supervision management, proposal-attempt recording, and panel coordination modules, the system enables precise oversight and automated administrative processes that optimize departmental operations while strengthening accountability. Built on a modular and scalable architecture, the FYP Tracking & Accountability System supports transparent management by adapting to varying operational scales and user requirements.")
empty()
body("Unified Modeling Language (UML) diagrams are used to visually represent system components, user interactions, and data flows, ensuring clear communication among developers, stakeholders, and end-users throughout the development process. The system’s design emphasizes intuitive user interfaces for all roles, ensuring accessibility for users with varying levels of technical expertise. Through its integrated approach, the platform provides a complete solution that transforms traditional final year project operations into modern, efficient, and transparent digital processes.")

h("Unified Modeling Language (UML)", 5)
body("In research projects involving the development and implementation of technological systems such as the FYP Tracking & Accountability System it is crucial to communicate system functionality in a clear and structured manner. One of the most effective ways to achieve this is through the use of Unified Modeling Language (UML). UML is a standardized visual modeling language that facilitates the design, documentation, and analysis of system components, user interactions, and data flows. It is widely adopted in software and systems engineering to provide a structured representation of how a system is constructed and operates.")

h("Benefits of Using UML in Software Development", 5)
body("In the development of a digital final year project management system for efficient academic operations, various stakeholders, system components, and processes must work together seamlessly to record, process, and respond to operational data. Unified Modeling Language (UML) plays a crucial role in this process by providing clear visual representations of the system architecture and workflows. First, UML diagrams help clarify project requirements and establish well-defined system boundaries, ensuring the final solution remains aligned with its core objectives of operational efficiency and accountability.")
empty()
body("Beyond planning, these diagrams serve as valuable technical documentation that demonstrates rigorous system design, adding credibility and professionalism to research publications and project reports. Additionally, UML acts as a universal communication tool that bridges understanding between developers, departmental staff, and project supervisors, enabling more productive discussions about system features, design choices, and implementation strategies.")
empty()
body("By standardizing how system elements like the state-tracking engine, supervision module, and notification system interact, UML facilitates efficient collaboration while reducing potential misunderstandings throughout the development lifecycle. The use of UML in the FYP Tracking & Accountability System project ensures that all stakeholders share a common understanding of the system’s structure and behavior, leading to more accurate implementation and easier maintenance of the final product.")

h("Use Case Diagram", 3)
empty()
body("Use case diagrams are visual tools within UML that illustrate the interactions between a system and its external environment, capturing the essential business requirements for system operation. These diagrams represent a business entity or software system, its external stakeholders (known as actors), and a set of tasks (use cases) that users are expected or authorized to perform when interacting with the system. They are particularly useful for defining the system’s functionality from the viewpoint of its users.")
empty()
body("By providing a clear, concise overview of how users interact with the system, use case diagrams help both developers and stakeholders better understand the system requirements. Additionally, they serve as an effective communication tool, allowing project teams to visually map out the roles and actions involved in achieving system goals.")
empty()
body("The four elements of a use case diagram are:")
empty()
body_label("System", "The FYP Tracking & Accountability System boundary")
body_label("Actors", "External entities interacting with the system (Student, Supervisor, Facilitator, HOD, Examiner, Superadmin)")
body_label("Use Cases", "Specific functionalities the system provides")
body_label("Relationships", "Connections between actors and use cases")
empty()
body("These diagrams utilize the following symbols:")
h("Actor", 5)
body("A stick figure symbol represents an external entity that interacts with the system. When directly interacting with a system, an external entity takes on a designated role defined by an actor. This role might represent a user’s function (such as Supervisor or Student) or a role fulfilled by another system that engages with the given system.")
h("Use case", 5)
body("The use case entails detailing the sequence of actions that a system can undertake while interacting with external actors. It encompasses tasks that the system should perform in response to an actor’s request.")
h("Relationship", 5)
body("Genuine associations illustrate the direct interactions between actors and use cases in a system. These are represented using the UML association symbol, indicating a meaningful connection within the system’s functionality. This helps ensure that all user interactions are clearly defined and properly linked to the system’s operations.")
h("System boundary", 5)
body("A box is drawn around the use case diagram to visually represent the system’s boundary. This defines the scope of the modeled system and distinguishes internal functionalities from external interactions. It helps stakeholders clearly identify what is included within the system and what lies outside its operational scope.")

h("Use Case Diagram of the System", 3)
empty()
body("The use case diagram represents the main interactions between users and the FYP Tracking & Accountability System. The system has six primary actors: Student, Supervisor, Facilitator, Head of Department (HOD), Examiner, and Superadmin, together with two external systems — the prototype-review system and the email provider.")
empty()
fig(2, "Use case diagram", 2)

h("Use Case Description", 5)
body("A Use Case description specifies what a use case accomplishes, and what it requires to be properly performed. Each use case is documented using the following structure:")
empty()
for item in ["Name: A unique identifier for the use case",
             "Description: A brief explanation of what the system aims to accomplish through this use case",
             "Actor: The primary participant(s) involved in the use case",
             "Precondition: The system state required before the use case may start",
             "Post-condition: The system state after the use case is successfully completed",
             "Normal Flow: The primary sequence of steps describing the standard scenario",
             "Alternative Flow: Variations or exception conditions that could occur during execution"]:
    body(item)

empty()
body("Table 1: Adding users Use Case Table", bold=True)
uc_table([
    ("Use Case Number", "UC-01"),
    ("Use Case Name", "User Authentication and Access"),
    ("Actor", "Student, Supervisor, Facilitator, HOD, Examiner, Superadmin"),
    ("Description", "Logging into the FYP Tracking & Accountability System and accessing role-specific functionality"),
    ("Pre-condition", "A user must have an account provisioned by the Superadmin, and students must have a verified email imported from the registrar's list"),
    ("Post-condition", "A user gains secure, token-based access to the system with appropriate role-based permissions"),
    ("Normal Flow", "Login: User enters credentials (email and password). System verifies credentials and issues a session token. View Profile: User can view and verify their account information. Access Dashboard: User views a role-specific dashboard. Perform Role Activities: Students view their state and timeline; supervisors post availability and confirm meetings; examiners record outcomes; HOD and facilitator coordinate the workflow. View Records: Access timelines, notifications, and relevant logs."),
    ("Alternative Flow", "If login credentials are invalid, the system displays an error message and allows retry. If the account is disabled, the system denies access and notifies the user. If a user attempts an action outside their role, the system rejects it. If the system encounters an error, the transaction is rolled back and the user is notified."),
])

h("Use case description for the HOD", 5)
body("Table 2: HOD Use Case Table", bold=True)
uc_table([
    ("Use Case Number", "UC-02"),
    ("Use Case Name", "Manage Student Journey"),
    ("Actor", "Head of Department (HOD)"),
    ("Description", "Managing student records and academic decisions within the FYP Tracking & Accountability System"),
    ("Pre-condition", "HOD must be logged into the system and must have the appropriate academic privileges"),
    ("Post-condition", "Student records, approvals, and assignments are correctly created or updated, and all changes are recorded in the audit trail"),
    ("Normal Flow", "Upload Student List: HOD imports the registered-student list with verified emails. Approve Case Letter: HOD reviews and approves case-study letters and issues preliminary requirements. Review Proposal: HOD records the accept or reject decision with a reason. Assign Supervisor: HOD assigns a supervisor to a student once the proposal is accepted. Read Prototype Data: HOD queries the external prototype system by student ID. Monitor: HOD views student states, timelines, and records."),
    ("Alternative Flow", "If an imported record is incomplete, the system highlights the missing fields and prevents submission. If an invalid state transition is attempted, the system rejects it and explains why. If a proposal has reached the maximum number of rejected attempts, the system locks resubmission until unlocked. If a database error occurs, changes are rolled back and the HOD is notified."),
])

h("Use case description for Backend side", 5)
body("Table 3: Backend side Table", bold=True)
uc_table([
    ("Use Case Number", "UC-03"),
    ("Use Case Name", "Automated State Transition and Notification"),
    ("Actor", "Backend (System)"),
    ("Description", "Automatically validating state transitions, counting attempts, and dispatching notifications based on defined business rules"),
    ("Pre-condition", "Student records are loaded and current; the state machine and notification templates are configured; the email queue is available"),
    ("Post-condition", "Valid transitions are applied and recorded; attempts are counted; notifications are queued, sent, and logged with delivery status"),
    ("Normal Flow", "Receive Action: System receives a transition or event request from an authorized actor. Validate Transition: The state machine confirms the requested transition is allowed from the current state. Apply Business Rules: Increment prototype or proposal attempt counters; enforce the maximum-attempt limit; keep panel assignments independent of supervision. Persist and Record: Update the student state, write a state-transition record, and append to the immutable audit log. Dispatch Notification: Render the milestone email template, queue it, send it asynchronously, and write to the notification log with delivery status."),
    ("Alternative Flow", "If the transition is invalid, the system raises an InvalidStateTransitionException and rejects the action. If the actor is not authorized, the action is denied and logged. If an email send fails, the notification is retried and the failure is recorded. If a database error occurs, the transaction is rolled back and the error is logged for technical review."),
])

h("Class Diagram", 3)
empty()
body("A class diagram serves as a fundamental component of software modeling and design in the field of software engineering. It provides a visual representation of the building blocks of the FYP Tracking & Accountability System, aiding in the understanding and planning of the software architecture. Within a class diagram, each class typically consists of three key elements: attributes, operations (methods), and associations. Attributes represent the properties or data members of a class, while operations define the behaviors or functions that the class can perform. Associations illustrate the relationships and connections between classes, allowing designers to depict how different classes interact within the system and influence one another’s behavior.")
empty()
body("In UML, class diagrams play a crucial role in modeling the static structure of the FYP Tracking & Accountability System. They not only reveal the relationships between classes but also showcase source code dependencies, inheritance hierarchies, and interface implementations. This makes class diagrams an indispensable tool for both high-level architectural design and low-level coding details, bridging the gap between abstract design and concrete implementation while ensuring consistency and traceability throughout the development process.")
empty()
body("Class diagrams are applicable to various object-oriented programming languages and paradigms, such as Java, C++, and Python. They offer a blueprint for creating, maintaining, and understanding the FYP Tracking & Accountability System by presenting a clear hierarchy of classes and their associations, improving both development efficiency and collaboration among diverse project team members and stakeholders across the software development lifecycle.")
empty()
body("The class diagram represents the static structure of the FYP Tracking & Accountability System, showing its main classes, attributes, and relationships.")
empty()
fig(3, "Class diagram", 3)

h("Sequence Diagram", 5)
body("A sequence diagram is a type of interaction diagram that depicts objects as lifelines running down the page, with messages rendered as arrows from the source lifeline to the target lifeline, representing their interactions through time. Object interactions are arranged in temporal sequence in a sequence diagram.")
empty()
body("In UML, the stages required to perform an operation are described in sequence diagrams, which are interaction diagrams. They show how objects interact when operating within a cooperative system. Sequence Diagrams are time-focused and use the vertical axis of the diagram to indicate time, what messages are received when, and how the interaction is organized graphically.")

h("The notations and their definitions that are used in sequence diagram", 5)
body("Table 4: Sequence diagram Table", bold=True)
tbl(
    ["Term and Definition", "Symbol"],
    [
        ["Actor: Could be a person or external system that interacts with but is not part of the system. Participates in interactions by sending and/or receiving messages. Positioned at the top of the diagram.", "Stick figure"],
        ["Object Lifeline: Represents an object that participates in a sequence by sending and/or receiving messages. Shows the lifespan of the object during the interaction. Placed vertically in the diagram with object name at top.", "Vertical dashed line"],
        ["Activation Box: A narrow rectangle positioned on top of a lifeline. Indicates when an object is active and performing an operation. Shows the period of message transmission and reception.", "Narrow rectangle"],
        ["Message: Synchronous Message: solid arrow with filled head → carries information and waits for response. Asynchronous Message: solid arrow with line head → sends message without waiting. Return Message: dashed arrow → indicates response or return value. Creation Message: dashed arrow with “new” → object instantiation. Destruction Message: solid arrow with X → object deletion.", "Arrows"],
    ],
    widths=[4.5, 2.0]
)

h("User Signup and Login Process", 5)
body("The FYP Tracking & Accountability System provides a secure and streamlined process for user authentication. Accounts are provisioned by the Superadmin, and students are imported from the registrar’s verified list. During login, users enter their credentials, which are verified against stored records secured with BCrypt hashing. Upon successful authentication, the system issues a JSON Web Token and grants access to role-specific functionalities.")
empty()
fig(4, "Sequence Diagram", 4)

h("Process Interpretation:", 5)
body("The authentication process begins with the user submitting login credentials, whose format is validated before the system checks them against stored records.")
empty()
body("If the credential format is incorrect, an error message is shown to the user. If the account is disabled or does not exist, a corresponding error message is displayed.")
empty()
body("For valid credentials, the password is verified against the stored BCrypt hash, a JSON Web Token is issued, and the system redirects the user to their role-specific dashboard.")
empty()
body("The session token is then included with each subsequent request, validating the user’s identity and role for every protected operation until the session expires.")

h("State transition sequence diagram:", 5)
fig(5, "State transition sequence diagram", 5)
body("The system begins with login and role validation, where user credentials are verified, and the actor’s role is determined.")
empty()
body("It proceeds to transition validation, ensuring the requested change is allowed from the student’s current state. If the transition is invalid, the system rejects it and raises an InvalidStateTransitionException.")
empty()
body("A business-rule check follows, confirming attempt counters and limits — for example, locking proposal resubmission after the third rejection. Error notifications are displayed if rules are violated.")
empty()
body("Once validations are complete, the system persists the new state, writes the state-transition record, and appends to the immutable audit log.")
empty()
body("The Notification Service then renders the appropriate milestone email, queues it for asynchronous delivery to the relevant recipient(s), and records it in the notification log.")
empty()
body("Finally, the database updates all relevant records, including the student state, transition history, and notification status for future tracking and dispute resolution.")

h("Activity Diagram", 5)
body("An activity diagram is a type of UML diagram that illustrates the workflow of the FYP Tracking & Accountability System processes. It shows the sequence of activities, decision points, and parallel processes, highlighting the flow of control from one task to another. Activity diagrams help visualize final year project management processes, clarify responsibilities among actors, and identify potential improvements by providing a detailed view of the system’s behavior.")

h("Key Elements in FYP Tracking & Accountability Context:", 5)
for item_h, item_b in [
    ("Activities", "Final year project management tasks or operations performed in the workflow. Notation: Rounded rectangles."),
    ("Actions", "A single step within an activity that cannot be broken down further. Notation: Rounded rectangles."),
    ("Initial Node (Start Node)", "The starting point of the project-management workflow. Notation: Filled black circle."),
    ("Final Node (End Node)", "The end point of the workflow. Notation: Filled black circle inside a larger unfilled circle."),
    ("Decision Nodes", "Points where the workflow branches based on conditions. Notation: Diamond shape."),
    ("Merge Nodes", "Points where multiple branches converge back into a single flow. Notation: Diamond shape."),
    ("Fork Nodes", "Points where a single flow splits into multiple concurrent flows. Notation: Thick horizontal or vertical bar with multiple outgoing arrows."),
    ("Join Nodes", "Points where multiple concurrent flows synchronize back into a single flow. Notation: Thick horizontal or vertical bar with multiple incoming arrows."),
    ("Swim lanes", "Partition the activity diagram to represent different actors (Student, Supervisor, HOD, System). Notation: Vertical or horizontal sections."),
    ("Transitions (Edges)", "Arrows that show the flow of control from one activity or action to another. Notation: Arrows connecting elements."),
    ("Signals and Events", "Represent asynchronous triggers or external inputs affecting the workflow. Notation: Special icons or symbols."),
]:
    body_label(item_h, item_b)
    empty()

h("User Account Creation and Login Process", 5)
fig(6, "User signup and login activity diagram", 6)

h("Student Journey Activity Diagram", 5)
fig(7, "Student journey activity diagram", 7)

h("Database Schema Diagram", 3)
empty()
body("A database schema diagram describes how data is organized to create a blueprint for how a database will be constructed and is the database management system’s supporting formal language used to define the structure of a database system (DBMS). Formally speaking, a database schema is a set of rules (sentences referred to as integrity constraints) applied to a database. The compatibility of the schema’s components is ensured by these integrity requirements.")
empty()
fig(8, "Database schema", 8)

h("Data Dictionary", 3)
empty()
body("A data dictionary serves as a vital component in the field of data management and database design. Its role extends beyond assisting programmers and includes supporting a wide range of stakeholders, from data analysts to database administrators, and even end-users. When designing or evaluating a system that involves user-interactive objects, the initial phase requires a comprehensive understanding of each object’s attributes and its relationships with other objects.")
empty()
body("This understanding helps in creating a clear and accurate data model that serves as the foundation for effective system design. A well-maintained data dictionary acts as a living document that evolves with the system. It serves as a reference point throughout the entire software development lifecycle, aiding in system maintenance, troubleshooting, and future enhancements. As systems grow in complexity, the data dictionary plays a pivotal role in managing data consistency and ensuring that all stakeholders share a common understanding of the data’s meaning and usage.")

h("User account", 5)
body("The User account Data Dictionary provides detailed descriptions of the attributes stored within the users database table.")
empty()
tbl(
    ["Field Name", "Data Type", "Description"],
    [
        ["id",               "UUID, Primary Key",     "Unique identifier for each user"],
        ["email",            "VARCHAR(255), Unique",  "User’s login email address"],
        ["full_name",        "VARCHAR(255)",          "User’s full name"],
        ["password_hash",    "VARCHAR(60)",           "BCrypt-hashed password for security"],
        ["phone",            "VARCHAR(20)",           "User’s contact number"],
        ["role",             "ENUM('STUDENT','SUPERVISOR','FACILITATOR','HOD','EXAMINER','SUPERADMIN')", "Role of the user in the system"],
        ["eligible_examiner","BOOLEAN, Default: FALSE","Indicates if the user can be assigned as an examiner"],
        ["enabled",          "BOOLEAN, Default: TRUE", "Indicates if the user account is active"],
        ["created_at",       "TIMESTAMP",             "Account creation timestamp"],
    ],
    widths=[1.8, 2.2, 2.5]
)

body("Students Table", bold=True)
tbl(
    ["Field Name", "Data Type", "Description"],
    [
        ["id",              "UUID, Primary Key",           "Unique identifier for each student"],
        ["user_id",         "UUID, Foreign Key → users(id)", "Links the student to a user account"],
        ["reg_number",      "VARCHAR(50), Unique",         "Student registration number"],
        ["organisation",    "VARCHAR(255)",                "Organisation or company for industry-based projects"],
        ["project_topic",   "VARCHAR(255)",                "The student’s project topic"],
        ["group_label",     "VARCHAR(100)",                "Group or cohort label"],
        ["state",           "ENUM (13 states)",            "Current state in the journey (e.g., SUPERVISION)"],
        ["state_entered_at","TIMESTAMP",                   "When the student entered the current state"],
        ["supervisor_id",   "UUID, Foreign Key → users(id)", "The assigned supervisor"],
        ["book_signed_off", "BOOLEAN, Default: FALSE",     "Whether the supervisor has signed off the book"],
        ["proto_attempts",  "INT, Default: 0",             "Number of prototype re-presentations"],
        ["proposal_locked", "BOOLEAN, Default: FALSE",     "Whether proposal resubmission is locked after three rejections"],
        ["flagged",         "BOOLEAN, Default: FALSE",     "Whether the student is flagged for attention"],
    ],
    widths=[1.8, 2.2, 2.5]
)

body("Meetings Table", bold=True)
tbl(
    ["Field Name", "Data Type", "Description"],
    [
        ["id",           "UUID, Primary Key",                   "Unique identifier for each meeting"],
        ["student_id",   "UUID, Foreign Key → students(id)",      "The student attending the meeting"],
        ["supervisor_id","UUID, Foreign Key → users(id)",         "The supervisor who scheduled the meeting"],
        ["scheduled_at", "TIMESTAMP",                           "The date and time of the meeting"],
        ["confirmed",    "BOOLEAN, Default: FALSE",             "Whether the supervisor confirmed the meeting"],
        ["attended",     "BOOLEAN, Nullable",                   "Attendance outcome once the meeting has passed"],
        ["topic",        "VARCHAR(255)",                        "The meeting topic"],
        ["notes",        "TEXT",                                "Notes recorded after the meeting"],
        ["meeting_type", "ENUM('ONLINE','IN_PERSON')",          "The mode of the meeting"],
        ["location",     "VARCHAR(255)",                        "Location or meeting link"],
    ],
    widths=[1.8, 2.2, 2.5]
)

body("Proposal Attempts Table", bold=True)
tbl(
    ["Field Name", "Data Type", "Description"],
    [
        ["id",               "UUID, Primary Key",               "Unique identifier for each proposal attempt"],
        ["student_id",       "UUID, Foreign Key → students(id)", "The student who submitted the proposal"],
        ["attempt_number",   "INT",                             "The sequential attempt number (1–3)"],
        ["status",           "ENUM('PENDING','ACCEPTED','REJECTED')", "The review outcome of the attempt"],
        ["rejection_reason", "TEXT",                            "The reason given when an attempt is rejected"],
        ["reviewed_by",      "UUID, Foreign Key → users(id)",    "The reviewer who recorded the decision"],
        ["submitted_at",     "TIMESTAMP",                       "When the attempt was submitted"],
        ["reviewed_at",      "TIMESTAMP",                       "When the attempt was reviewed"],
    ],
    widths=[1.8, 2.2, 2.5]
)

body("Panel Assignments Table", bold=True)
tbl(
    ["Field Name", "Data Type", "Description"],
    [
        ["id",                  "UUID, Primary Key",               "Unique identifier for each panel assignment"],
        ["student_id",          "UUID, Foreign Key → students(id)", "The student being examined"],
        ["examiner_id",         "UUID, Foreign Key → users(id)",    "The assigned examiner"],
        ["panel_type",          "ENUM('PRE_DEFENSE','DEFENSE')",   "The type of panel"],
        ["scheduled_at",        "TIMESTAMP",                       "The scheduled date and time"],
        ["outcome",             "ENUM('CLEARED','PASSED','REFERRED','FAILED')", "The recorded panel outcome"],
        ["outcome_note",        "TEXT",                            "Notes recorded with the outcome"],
        ["assigned_by",         "UUID, Foreign Key → users(id)",    "The facilitator who made the assignment"],
        ["assigned_at",         "TIMESTAMP",                       "When the assignment was made"],
    ],
    widths=[1.8, 2.2, 2.5]
)

h("System architecture diagram", 3)
empty()
body("The system architecture of the FYP Tracking & Accountability System illustrates how the platform’s components interact to deliver a seamless and efficient final year project management experience.")
empty()
fig(9, "System architecture", 9)

# ═══════════════════════════════════════════════════════════════════════════════
# CHAPTER 4
# ═══════════════════════════════════════════════════════════════════════════════
pb()
h("CHAPTER 4", 2)
pch4 = doc.add_paragraph(style='Heading 2')
_run(pch4, "IMPLEMENTATION OF THE FYP TRACKING & ACCOUNTABILITY SYSTEM", 16, bold=True)

h("Introduction", 3)
empty()
body("This chapter presents the implementation of the FYP Tracking & Accountability System, a web-based platform developed to modernize and document the management of the final year project journey at the Adventist University of Central Africa. It translates the requirement analysis and design discussed in the previous chapter into a functioning system, describing the technologies adopted, the structure of the application, and the way each module was realized in practice.")
empty()
body("The implementation phase focused on digitizing the existing manual workflow, including student registration, case-letter approval, prototype review, proposal-attempt tracking, supervisor assignment, supervision meetings, panel assignment, and final defense. Particular attention was given to recording every state transition with the actor who performed it and the time it occurred, so that the project journey becomes a continuous, timestamped, and auditable record rather than a collection of scattered documents.")
empty()
body("The chapter also discusses the integration of the user-interface components, the backend services, the relational database, and the supporting infrastructure such as caching, authentication, and email notification. Each major screen of the system is presented to demonstrate how the design objectives were achieved and how the different roles — student, supervisor, facilitator, Head of Department, examiner, and superadmin — interact with the platform.")
empty()
body("Finally, the chapter presents the software testing techniques applied to verify the correctness and reliability of the system, followed by the hardware and software specifications required for deployment in a production environment.")

h("Technologies Used", 3)
empty()
body("To build a robust, secure, and responsive platform for the FYP Tracking & Accountability System, a combination of modern front-end, back-end, and database technologies was employed. These technologies were selected for their reliability, scalability, strong community support, and suitability for building an auditable, role-based academic management system.")
h("Front-End Development:", 5)
body_label("React 19 with TypeScript", "The user interface was built using React, a component-based JavaScript library, together with TypeScript for static type checking. This combination enabled the creation of dynamic, reusable, and type-safe interface components for the different user roles, reducing runtime errors and improving maintainability.")
empty()
body_label("Vite", "Vite was used as the front-end build tool and development server, providing fast module bundling, instant hot-reloading during development, and optimized production builds.")
empty()
body_label("CSS and Responsive Design", "Styling and layout were implemented with modern CSS to ensure that the dashboards, forms, and tables render correctly on both desktop and mobile-sized screens, allowing staff and students to access the system from a range of devices.")
empty()
body_label("Axios", "The Axios HTTP client was used to communicate with the backend REST API, attaching the JSON Web Token to each authenticated request and handling responses and errors consistently across the application.")
empty()
body_label("Visual Studio Code", "VS Code served as the primary integrated development environment for the front end, offering rich TypeScript support, debugging tools, and extensions that streamlined development.")
h("Back-End & Database Development:", 5)
body_label("Java 17 and Spring Boot", "The server-side logic was developed using Java 17 and the Spring Boot 3.2.5 framework, which provided a production-ready foundation for building REST APIs, dependency injection, and configuration management with minimal boilerplate.")
empty()
body_label("Spring Security and JWT", "Authentication and authorization were implemented with Spring Security and JSON Web Tokens (jjwt). Passwords are hashed using BCrypt, and a stateless security filter validates the bearer token on every request, enforcing role-based access control across all endpoints.")
empty()
body_label("Spring Data JPA and Hibernate", "Data access was implemented using Spring Data JPA with Hibernate as the object-relational mapping provider, allowing entities such as User, Student, StateTransition, Meeting, ProposalAttempt, and PanelAssignment to be persisted and queried in an object-oriented manner.")
empty()
body_label("PostgreSQL and Flyway", "PostgreSQL 16 was used as the relational database for secure and reliable storage of all records. Flyway managed database schema versioning through migration scripts that run automatically on application startup, ensuring a consistent and reproducible database structure.")
empty()
body_label("Redis", "Redis was integrated as an in-memory cache to improve performance for frequently accessed data and to support fast, scalable read operations.")
empty()
body_label("Apache POI", "The Apache POI library was used to implement bulk student import from Excel spreadsheets, allowing the department to register an entire cohort from a single file rather than entering students one by one.")
empty()
body_label("Spring Mail", "Email notifications were implemented using Spring Mail (JavaMailSender), which sends milestone emails to students, supervisors, and examiners, with every message recorded in a notification log.")
empty()
body_label("Maven", "Apache Maven was used for build automation and dependency management of the backend project.")

h("Presentation of the New System", 3)
empty()
body("This section illustrates the results of the implementation of the FYP Tracking & Accountability System. The following figures present the main interfaces of the platform and explain the role each one plays in the management of the final year project journey.")

h("Login Page:", 5)
body("A secure login interface through which all users — students, supervisors, facilitators, the Head of Department, examiners, and the superadmin — authenticate before accessing the system. Credentials are verified against the stored BCrypt hash, and on success a JSON Web Token is issued and used to authorize subsequent requests.")
placeholder_fig(10, "Login Page", "the login screen showing the email and password fields, the AUCA branding, and the sign-in button")

h("Dashboard:", 5)
body("A role-aware dashboard that gives the Head of Department and facilitators an overview of the cohort, including the number of students at each state of the project journey and quick access to the main management functions.")
placeholder_fig(11, "Dashboard", "the main dashboard showing summary cards for each student state and navigation to students, supervision, panels, and logs")

h("User Management:", 5)
body("An interface that allows the superadmin to create user accounts, assign roles, enable or disable accounts, and reset passwords, ensuring controlled access to the system.")
placeholder_fig(12, "User Management Page", "the user management screen listing users with their roles and enable/disable and reset-password actions")

h("Student Registration and Excel Import:", 5)
body("An interface that allows the department to register students individually or to import an entire cohort from an Excel file. The import reads the reg number, full name, email, phone, organization, and group columns and creates the corresponding student records automatically.")
placeholder_fig(13, "Student Import Page", "the student import screen with the file-upload control and a preview of imported students")

h("Student Profile and State Timeline:", 5)
body("A detailed view of a single student showing their current state, assigned supervisor, and the complete timeline of state transitions. Each entry records the previous and new state, the actor who performed the change, and the timestamp, providing a continuous and auditable history.")
placeholder_fig(14, "Student Profile and State Timeline", "a student detail page showing the state timeline, supervisor details, and transition buttons")

h("Supervision Module:", 5)
body("An interface through which supervisors post their weekly availability slots, schedule and confirm meetings with their assigned students, and record meeting attendance and notes. This closes the supervision blind spot identified in the analysis of the existing system.")
placeholder_fig(15, "Supervision Page", "the supervision screen showing availability slots and a list of scheduled meetings with confirm and outcome actions")

h("Panel Assignment:", 5)
body("An interface used by the facilitator to assign examiners to pre-defense and defense panels and to record outcomes. The system enforces the rule that a student’s own supervisor cannot be assigned as their examiner.")
placeholder_fig(16, "Panel Assignment Page", "the panel assignment screen showing examiners assigned to a student’s pre-defense and defense panels")

h("Notification and Audit Logs:", 5)
body("Oversight views that display every email the system has sent, with its delivery status, and an immutable audit log of actions performed in the system. These records allow the department to resolve disputes from evidence rather than assumption.")
placeholder_fig(17, "Notification and Audit Logs", "the notification log listing sent emails with statuses, alongside the audit log of recorded actions")

h("Software Testing", 3)
empty()
body("The design of software benefits from the use of software tests. They aid in determining whether the system meets its requirements and behaves as intended before it is deployed for real use. Testing was carried out progressively throughout the implementation of the FYP Tracking & Accountability System to detect and correct defects early.")
empty()
body("When undertaking software testing, it is vital to keep the following questions in mind:")
empty()
body("Is the application compliant with the standards that guided its design and development?")
body("Is the application functioning as it should?")
body("Is it possible to deploy the application reliably and satisfy the needs of the stakeholders?")
empty()
body("The following software testing techniques were applied:")
empty()
body_label("Unit Testing", "Unit testing was used to verify the correct operation of individual components in isolation, particularly the state machine that governs valid transitions between the thirteen states of the student journey. Tests confirmed that valid transitions are accepted and that invalid transitions raise the appropriate exception.")
empty()
body_label("Integration Testing", "Integration testing verified that the combined components — controllers, services, repositories, the database, and the security filter — work together correctly. The REST endpoints were exercised against a running server to confirm authentication, role enforcement, student listing, Excel import, and supervision operations behaved as expected.")
empty()
body_label("Validation Testing", "Validation testing confirmed that the system behaves correctly from the user’s perspective and that input is properly validated. For example, the email address must be well formed, required fields must be supplied, and only authorized roles may perform protected actions; otherwise a clear error message is returned.")

h("Hardware and Software Requirements", 3)
empty()
body("Client-Side Software Requirements:", bold=True)
for item in [
    "A modern web browser (Google Chrome, Mozilla Firefox, Microsoft Edge, or Opera).",
    "Operating system: Windows 10/11, Linux, or macOS.",
    "A RAM of 4 GB (minimum).",
    "A hard disk with at least 2 GB of free space.",
    "A stable internet connection to reach the application server.",
]:
    body(item)
empty()
body("Server-Side Software Requirements:", bold=True)
for item in [
    "Operating system: a Linux server distribution (or Windows Server) capable of running the Java runtime.",
    "Java 17 runtime environment for the Spring Boot application.",
    "PostgreSQL 16 database server.",
    "Redis server for caching.",
    "RAM: 8 GB (minimum) recommended for the application, database, and cache.",
    "Network: 1 Gbps network interface for reliable access under concurrent load.",
]:
    body(item)

# ═══════════════════════════════════════════════════════════════════════════════
# CHAPTER 5
# ═══════════════════════════════════════════════════════════════════════════════
pb()
h("CHAPTER 5", 2)
pch5 = doc.add_paragraph(style='Heading 2')
_run(pch5, "CONCLUSION AND RECOMMENDATIONS", 16, bold=True)

h("Introduction", 3)
empty()
body("This chapter presents the conclusion of the research and development of the FYP Tracking & Accountability System. It summarizes the work accomplished throughout the five chapters of this study, evaluates the extent to which each specific objective was achieved, reflects on the limitations encountered during development, and proposes recommendations for future enhancements. The chapter concludes with a final statement on the significance of the system and its potential impact on final year project administration at the Adventist University of Central Africa (AUCA).")

h("Summary of the Work", 3)
empty()
body("This study set out to address a clearly identified problem: the management of final year projects at AUCA relied on fragmented, manual, and paper-based processes that created inefficiencies, reduced transparency, and made accountability difficult to enforce. Students were uncertain about their current stage, supervision events were unverifiable, proposal rejection history was anecdotal rather than documented, and panel coordination depended on informal communication. The department lacked a single authoritative source of truth for any student’s complete journey.")
empty()
body("In response, the FYP Tracking & Accountability System was designed and developed as a full-stack, role-based web application that digitizes and enforces every stage of the final year project journey. The system implements a thirteen-state machine covering the complete lifecycle from registration through case-letter submission and approval, prototype review, proposal submission, supervision, book submission, pre-defense, defense, and final completion or withdrawal. Six distinct user roles — Student, Supervisor, Facilitator, Head of Department, Examiner, and Superadmin — interact with the system according to strictly enforced, role-based access rules.")
empty()
body("The backend was built using Spring Boot 3.2.5 and Java 17, with Spring Security providing JWT-based authentication and authorization. PostgreSQL 16 serves as the relational database, managed through Flyway schema migrations, while Redis 7 provides JWT blacklisting and caching support. The frontend was developed using React 19 and TypeScript, delivering responsive, role-specific dashboards that are wired entirely to the live backend API. The complete system was tested through a suite of 45 passing tests, including unit tests for the state machine and proposal service, and integration tests that verified the full student journey from registration to completion against a real PostgreSQL database.")

h("Achievement of Objectives", 3)

h("Specific Objective 1: Secure, Web-Based, and Responsive Platform", 5)
body("The FYP Tracking & Accountability System was delivered as a fully functional web application accessible through any modern browser. The React 19 frontend provides responsive interfaces for all six user roles, and the Spring Boot backend exposes a well-structured REST API that processes all state transitions, supervision events, proposals, and panel assignments in real time. Manual and paper-based processes for registration, approval, and coordination have been entirely replaced by digital workflows, fulfilling this objective.")

h("Specific Objective 2: Intuitive and User-Friendly Interface for All Roles", 5)
body("Role-specific dashboards were developed for every actor in the system. Students see their current state, a visual timeline of every transition, their supervision schedule, proposal history with rejection reasons and remaining attempts, and panel assignments. Supervisors manage their availability slots, confirm meetings, record attendance notes, and sign off the student’s book. The HOD and Facilitator access a full cohort view with filtering by state, and can approve letters, assign supervisors, review proposals, assign examiners, and record panel outcomes. The Examiner sees only their assigned panels and can record outcomes. The Superadmin can manage all users, reset passwords, and monitor system-wide logs. Each interface presents only the actions relevant to the logged-in role, reducing complexity and supporting ease of use.")

h("Specific Objective 3: Robust Database System", 5)
body("A ten-table PostgreSQL schema was designed and implemented, covering users, students, state transitions, availability slots, meetings, proposal attempts, panel assignments, notification logs, audit logs, and WhatsApp groups. All relationships are enforced with foreign key constraints, all identifiers use UUIDs generated by the database, and all timestamps are stored with timezone awareness. Flyway manages schema migrations in version order, ensuring that the database evolves safely across environments. The data model captures every event in the student journey with enough detail to reconstruct the full history of any student at any point.")

h("Specific Objective 4: State Machine and Role-Based Business Rules", 5)
body("The thirteen-state machine is fully enforced in the StudentStateService, which validates every requested transition against a predefined map of allowed transitions before applying it. Invalid transitions throw a typed InvalidStateTransitionException that is handled globally and returned to the caller with a descriptive error. Business rules beyond simple state validation are also enforced: prototype re-presentation attempts are counted automatically; proposal submissions are tracked across attempts, rejection reasons are recorded, and submission is locked after three rejections until the HOD explicitly unlocks it; panel examiners cannot be the student’s own supervisor; and the WITHDRAWN state is terminal, preventing any further transitions once set. These rules operate entirely at the service layer, meaning they cannot be bypassed through the API.")

h("Specific Objective 5: Automated Communication and Notification System", 5)
body("The notification service sends milestone emails at every significant event in the student journey, including case-letter approval, supervisor assignment, proposal submission confirmation, proposal acceptance or rejection with reason and remaining attempts, proposal locking, proposal unlocking, and panel outcome. Every email sent or attempted is recorded in the notification_logs table with the template key, recipient, delivery status, error message if applicable, retry count, and timestamp. This gives the department verifiable proof of every communication sent, addressing one of the core complaints identified during the analysis phase.")

h("Specific Objective 6: Oversight and Reporting Tools", 5)
body("The HOD and Facilitator dashboards provide a real-time view of every student’s current state, with filtering by state to identify students requiring attention. Per-student detail pages show the full transition timeline, all supervision meetings with attendance records, the complete proposal attempt history with rejection reasons, and panel assignment outcomes. The immutable audit log records every action taken in the system — who performed it, in what role, on which entity, and at what time — providing the department with a reliable record for dispute resolution, compliance review, and evidence-based planning.")

h("Specific Objective 7: Data Security, Privacy, and System Reliability", 5)
body("Security was implemented in depth across all layers of the system. Authentication uses JWT tokens signed with HMAC-SHA384, with a configurable expiry of 24 hours. Logout blacklists the token in Redis so it cannot be reused. Passwords are hashed with BCrypt at cost factor 12, making brute-force attacks computationally expensive. Login attempts are rate-limited to five per IP address per fifteen minutes using Bucket4j, protecting against credential-stuffing attacks. All endpoints are protected by Spring Security’s role-based access control using @PreAuthorize annotations, ensuring that no user can perform actions outside their authorized role. HTTP security headers including HSTS, X-Frame-Options, X-Content-Type-Options, and XSS Protection are applied globally. Input validation using Bean Validation prevents malformed data from reaching the service layer.")

h("Limitations", 3)
empty()
body("While the FYP Tracking & Accountability System successfully achieves all seven specific objectives, a number of limitations were identified during development and testing that represent the boundaries of the current implementation.")
empty()
body_label("Absence of a Native Mobile Application", "The system is delivered as a responsive web application, which means it can be accessed from a smartphone browser. However, it does not currently provide a dedicated native application for Android or iOS. For students and supervisors who primarily use mobile devices, a native application would offer improved performance, push notifications, and offline access to key information such as the current state and upcoming meeting schedules.")
empty()
body_label("Email Delivery Dependency", "Automated email notifications require a configured SMTP server and valid credentials to function. In environments where these credentials are unavailable or where institutional email policies restrict outbound SMTP access, notifications will fail to deliver. While all delivery attempts are logged with their status, students and supervisors will not receive email alerts until the configuration is in place. The system currently has no fallback channel such as SMS or in-app push notification.")
empty()
body_label("Prototype Review System Integration is Stubbed", "The architecture includes an adapter for a read-only Prototype Review System API that would allow the HOD to retrieve a student’s prototype assessment data directly within the platform. At the time of completion, this integration is implemented as a configurable stub, returning no data until a real API endpoint is provided. The prototype submission and scoring process still takes place outside the system.")
empty()
body_label("No Real-Time Updates", "The frontend does not currently use WebSockets or server-sent events to push updates to connected clients in real time. Users must refresh their browser to see the latest state of students or meetings. For a department managing a large cohort during an active review period, real-time updates would significantly improve coordination and reduce the risk of acting on stale information.")
empty()
body_label("Limited Reporting and Analytics", "While the system provides a per-student timeline and audit log, it does not currently generate printable or exportable summary reports for the cohort as a whole. The HOD cannot currently export, for example, a PDF list of all students in the PRE_DEFENSE state, a statistical summary of average proposal attempts per cohort, or a supervision compliance report. These reports would support institutional oversight and planning beyond what the current dashboards provide.")
empty()
body_label("No Document Storage", "The system records that a student submitted a proposal or that a book was signed off, but it does not store the actual document files. Students cannot upload their proposal PDF or final book through the platform, and reviewers must share and review documents through separate channels. This limits the system’s ability to serve as a complete, end-to-end record of each student’s submitted work.")
empty()
body_label("WhatsApp Integration is Link-Only", "The WhatsApp group module stores the link and group name for each supervisor’s cohort, but does not integrate with the WhatsApp Business API to send automated messages directly. Notifications are sent only by email, and students who are more reachable through WhatsApp must still be contacted manually through those groups.")

h("Recommendations", 3)
empty()
body("Based on the limitations identified above and the operational needs observed during development, the following enhancements are recommended for future versions of the FYP Tracking & Accountability System.")
empty()
body_label("1. Develop a Native Mobile Application", "A dedicated mobile application for Android and iOS would allow students and supervisors to receive push notifications, check state and meeting status offline, and interact with the system more conveniently from their mobile devices. The existing REST API is already structured to support a mobile frontend without backend changes.")
empty()
body_label("2. Add Real-Time Communication Channels", "Integration of WebSocket-based real-time updates would allow dashboards to refresh automatically when a student’s state changes or a meeting is confirmed. In addition, SMS integration through a service such as Africa’s Talking would provide a reliable fallback notification channel for users without stable email access, which is particularly relevant in the Rwandan institutional context.")
empty()
body_label("3. Integrate with the Real Prototype Review System", "Once the prototype review API is documented and available, the existing adapter in the system should be configured to query it with the student’s registration number and display the assessment result directly in the student’s profile. This would eliminate the need for the HOD to consult a separate system when making the prototype-granted decision.")
empty()
body_label("4. Build a Cohort Analytics and Reporting Module", "A reporting module that allows the HOD and Facilitator to generate, filter, and export summary reports would significantly extend the system’s utility for institutional oversight. Recommended reports include: cohort state distribution at any given date, average proposal attempt counts per cohort cycle, supervision compliance summary by supervisor, and panel outcome statistics. These reports could be generated as downloadable PDF or Excel files from within the dashboard.")
empty()
body_label("5. Add Document Upload and Storage", "Future versions should allow students to upload their proposal book and final report directly through the platform. These uploads would be stored securely and linked to the corresponding proposal attempt or book submission event. Reviewers would then be able to access and annotate the submitted document from within the same interface, creating a fully self-contained record for each student.")
empty()
body_label("6. Integrate WhatsApp Business API for Direct Messaging", "Replacing the current link-only WhatsApp group module with a WhatsApp Business API integration would allow the system to send automated messages directly to a student’s WhatsApp number at key milestones, such as supervisor assignment, meeting confirmation, or panel scheduling. Given the high WhatsApp penetration among students and academic staff in Rwanda, this would substantially improve notification reach and engagement.")
empty()
body_label("7. Introduce Multi-Cohort Management", "The current system manages a single active cohort at a time. A multi-cohort architecture would allow the department to maintain records for multiple academic years simultaneously, compare cohort performance over time, and archive historical data without disrupting the active cycle. This is particularly important as the system is adopted across multiple departments or academic cycles.")
empty()
body_label("8. Expand Role Coverage with a Student Self-Registration Flow", "At present, students are imported by the HOD via Excel upload. A future enhancement could allow students to self-register using their institutional email address, with the HOD approving registrations before they are activated. This would reduce the administrative burden of maintaining and importing the registration spreadsheet at the start of each cycle.")

h("Conclusion", 3)
empty()
body("The FYP Tracking & Accountability System was developed to address a well-documented and operationally significant problem: the management of final year projects at AUCA through fragmented manual processes that undermined efficiency, transparency, and accountability. Through the systematic application of software engineering principles — including requirements analysis, database design, state machine modeling, role-based access control, and full-stack web development — a complete, functional, and tested digital platform was designed and delivered.")
empty()
body("All seven specific objectives defined at the outset of this study were achieved. The system enforces a thirteen-state student journey through a validated and tested state machine; provides six distinct, role-appropriate interfaces built on a live backend API; stores every event, decision, notification, and administrative action in a relational database with full audit capability; and protects all user data through industry-standard authentication, authorization, and encryption mechanisms.")
empty()
body("The platform is ready for deployment to a real cohort. With configuration of a production server, a managed database, and valid SMTP credentials, the FYP Tracking & Accountability System can serve the AUCA Final Year Project Department immediately, reducing the administrative workload that currently falls on the Head of Department and facilitators, providing students with real-time visibility of their own progress, and giving the institution an immutable, timestamped record of every decision made throughout the academic cycle.")
empty()
body("The limitations identified in this chapter do not undermine the system’s core functionality; they represent natural boundaries of a first-cycle implementation and a clear roadmap for future development. Each limitation has a corresponding recommendation that builds on the existing architecture without requiring fundamental redesign, demonstrating that the system was engineered with extensibility in mind.")
empty()
body("In conclusion, the FYP Tracking & Accountability System demonstrates that targeted digital transformation — even within a constrained academic development context — can produce tools that are immediately useful, technically sound, and aligned with real institutional needs. It is hoped that this work will serve not only the AUCA Final Year Project Department but also as a reference for similar institutions seeking to modernize their academic project management processes.")

# ── Save ──────────────────────────────────────────────────────────────────────
doc.save(OUTPUT)
print(f"Saved: {OUTPUT}")
